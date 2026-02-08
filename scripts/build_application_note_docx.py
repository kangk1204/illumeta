#!/usr/bin/env python3
"""
Build Bioinformatics Application Note DOCX for IlluMeta.
Generates a publication-ready ~4-page A4 document with:
- Title, Abstract, Introduction, Methods, Results, Discussion
- Table 1 (Feature comparison), Table 2 (Benchmark summary)
- Figure 1 (Workflow), Figure 2 (Benchmark multi-panel)
- Numbered references with PMIDs
"""
import argparse
import csv
import json
import os
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
except ImportError:
    sys.exit(
        "Missing python-docx.\n"
        "Install (inside your conda env): python -m pip install -r requirements-paper.txt\n"
        "Or (conda-forge): conda install -c conda-forge python-docx"
    )


# ── Helper utilities ──────────────────────────────────────────────────────────

def find_results_dir(gse_dir):
    """Find the results directory inside a GSE folder (handles varied naming)."""
    gse_path = Path(gse_dir)
    std = gse_path / 'illumeta_results'
    if std.exists() and (std / 'summary.json').exists():
        return std
    for d in sorted(gse_path.iterdir()):
        if d.is_dir() and d.name.endswith('_results') and (d / 'summary.json').exists():
            return d
    return None


def safe_float(val, default=0.0):
    try:
        if val is None or val == '' or val == 'NA':
            return default
        return float(val)
    except (ValueError, TypeError):
        return default


def safe_int(val, default=0):
    try:
        if val is None or val == '' or val == 'NA':
            return default
        return int(float(val))
    except (ValueError, TypeError):
        return default


def count_fdr_significant_csv_rows(csv_path, threshold=0.05):
    """Count rows with FDR/q-value < threshold for common column names."""
    fdr_cols = ['p.adjust', 'FDR', 'adj.P.Val', 'qvalue', 'q.value']
    try:
        with open(csv_path, 'r') as f:
            reader = csv.DictReader(f)
            if reader.fieldnames is None:
                return 0
            cols = set(reader.fieldnames)
            chosen = next((c for c in fdr_cols if c in cols), None)
            if chosen is None:
                return 0
            sig = 0
            for r in reader:
                try:
                    v = r.get(chosen, '')
                    if v is None or v == '' or v == 'NA':
                        continue
                    if float(v) < threshold:
                        sig += 1
                except (ValueError, TypeError):
                    continue
            return sig
    except OSError:
        return 0


def get_illumeta_version():
    """Best-effort: read __version__ from illumeta.py to stamp the manuscript."""
    illumeta_py = Path(__file__).resolve().parents[1] / 'illumeta.py'
    try:
        text = illumeta_py.read_text(encoding='utf-8')
    except OSError:
        return None
    m = re.search(r'__version__\s*=\s*[\"\']([^\"\']+)[\"\']', text)
    return m.group(1).strip() if m else None


def collect_benchmark_data(results_dir):
    """Collect benchmark metrics from results directories."""
    summary = []
    results_path = Path(results_dir)

    for gse_dir in sorted(results_path.glob('GSE*')):
        if not gse_dir.is_dir():
            continue
        res_dir = find_results_dir(gse_dir)
        if res_dir is None:
            continue

        gse_id = gse_dir.name
        row = {'gse_id': gse_id}

        # summary.json
        sj_file = res_dir / 'summary.json'
        if sj_file.exists():
            with open(sj_file, 'r') as f:
                sj = json.load(f)
                row['n_con'] = sj.get('n_con', 0)
                row['n_test'] = sj.get('n_test', 0)
                row['caf_score'] = sj.get('primary_caf_score', 0)
                row['crf_tier'] = sj.get('crf_sample_tier', '')
                n_total = sj.get('n_con', 0) + sj.get('n_test', 0)
                row['n_samples'] = str(n_total)
                row['consensus_dmps'] = (sj.get('intersect_up', 0) + sj.get('intersect_down', 0))

        # CRF_Sample_Tier.csv (tier only; n_samples uses post-QC count from summary.json)
        tier_file = res_dir / 'CRF_Sample_Tier.csv'
        if tier_file.exists():
            with open(tier_file, 'r') as f:
                reader = csv.DictReader(f)
                for r in reader:
                    row['crf_tier'] = r.get('tier', row.get('crf_tier', ''))
                    break

        # Platform from Probe_Filter_Summary.csv (raw probe count)
        probe_file = res_dir / 'Probe_Filter_Summary.csv'
        if probe_file.exists():
            with open(probe_file, 'r') as f:
                reader = csv.DictReader(f)
                for r in reader:
                    if r.get('step') == 'raw':
                        n_probes = safe_int(r.get('remaining', 0))
                        if n_probes > 900000:
                            row['platform'] = 'EPICv2'
                        elif n_probes > 600000:
                            row['platform'] = 'EPIC'
                        else:
                            row['platform'] = '450k'
                        break

        # Lambda from ablation
        for pipeline in ['Minfi', 'Sesame']:
            abl_file = res_dir / f'{pipeline}_Ablation_Summary.csv'
            if abl_file.exists():
                with open(abl_file, 'r') as f:
                    reader = csv.DictReader(f)
                    for r in reader:
                        if r.get('metric') == 'lambda':
                            row[f'lambda_{pipeline.lower()}'] = safe_float(
                                r.get('corrected', r.get('raw', '1.0')), 1.0)
                            row[f'lambda_{pipeline.lower()}_raw'] = safe_float(
                                r.get('raw', '1.0'), 1.0)
                            break

        # DMP counts (both FDR<0.05 and nominal: raw P<0.05 + |logFC|>1)
        for pipeline in ['Minfi', 'Sesame']:
            dmp_file = res_dir / f'{pipeline}_DMPs_full.csv'
            if dmp_file.exists():
                fdr_count = 0
                nominal_count = 0
                with open(dmp_file, 'r') as f:
                    reader = csv.DictReader(f)
                    for r in reader:
                        try:
                            fdr = float(r.get('adj.P.Val', r.get('FDR', '1')))
                            if fdr < 0.05:
                                fdr_count += 1
                            pval = float(r.get('P.Value', r.get('pval', '1')))
                            logfc = abs(float(r.get('logFC', '0')))
                            if pval < 0.05 and logfc > 1:
                                nominal_count += 1
                        except (ValueError, TypeError):
                            pass
                row[f'{pipeline.lower()}_dmps'] = fdr_count
                row[f'{pipeline.lower()}_dmps_nominal'] = nominal_count

        # Consensus DMPs from file (overrides summary.json)
        consensus_file = res_dir / 'Intersection_Consensus_DMPs.csv'
        if consensus_file.exists():
            with open(consensus_file, 'r') as f:
                row['consensus_dmps'] = sum(1 for _ in f) - 1

        # Count DMRs per pipeline
        for pipeline in ['Minfi', 'Sesame']:
            dmr_file = res_dir / f'{pipeline}_DMRs.csv'
            if dmr_file.exists():
                # DMR output includes non-significant regions; count only FDR<0.05.
                row[f'{pipeline.lower()}_dmrs'] = count_fdr_significant_csv_rows(dmr_file, threshold=0.05)

        if row.get('n_samples') and row.get('n_samples') != '0':
            summary.append(row)
            print(f"  Collected: {gse_id} (n={row['n_samples']}, tier={row.get('crf_tier', 'N/A')})")

    return summary


# ── Style helpers ─────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set table cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex
    })
    shading.append(shading_elem)


def add_styled_paragraph(doc, text, style_name='Normal', bold=False, italic=False,
                         font_size=9, alignment=None, space_after=Pt(4), space_before=Pt(0)):
    """Add a paragraph with inline formatting."""
    para = doc.add_paragraph()
    para.style = doc.styles[style_name] if style_name in [s.name for s in doc.styles] else doc.styles['Normal']
    run = para.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Arial'
    run.bold = bold
    run.italic = italic
    para.paragraph_format.space_after = space_after
    para.paragraph_format.space_before = space_before
    if alignment:
        para.alignment = alignment
    return para


def add_heading_styled(doc, text, level=1, font_size=10):
    """Add a section heading in Application Note style."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Arial'
    run.bold = True
    if level == 1:
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(4)
    else:
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(2)
    return para


def add_body_text(doc, text, font_size=9):
    """Add body text paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Arial'
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.space_before = Pt(0)
    # Set line spacing to single
    para.paragraph_format.line_spacing = Pt(12)
    return para


def add_body_with_refs(doc, text, font_size=9):
    """Add body text with superscript reference numbers.
    References are marked as [N] in the text and rendered as superscript."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.line_spacing = Pt(12)

    import re
    parts = re.split(r'(\[\d+(?:,\s*\d+)*\])', text)
    for part in parts:
        if re.match(r'\[\d+(?:,\s*\d+)*\]', part):
            run = para.add_run(part)
            run.font.size = Pt(7)
            run.font.name = 'Arial'
            run.font.superscript = True
        else:
            run = para.add_run(part)
            run.font.size = Pt(font_size)
            run.font.name = 'Arial'
    return para


# ── Document builder ──────────────────────────────────────────────────────────

def build_docx(benchmark_data, figures_dir, out_path, *, contact_email, repo_url, supplementary_text):
    """Build the complete Application Note DOCX."""
    doc = Document()

    # ── Page setup: A4, 25.4mm margins ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9)

    # ══════════════════════════════════════════════════════════════════════════
    # TITLE
    # ══════════════════════════════════════════════════════════════════════════
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(
        'IlluMeta: an automated dual-pipeline framework for robust DNA methylation analysis'
    )
    title_run.font.size = Pt(14)
    title_run.font.name = 'Arial'
    title_run.bold = True
    title_para.paragraph_format.space_after = Pt(6)

    # Authors
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_para.add_run('Keunsoo Kang')
    author_run.font.size = Pt(10)
    author_run.font.name = 'Arial'
    author_para.paragraph_format.space_after = Pt(2)

    # Affiliation
    affil_para = doc.add_paragraph()
    affil_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affil_run = affil_para.add_run('Department of Microbiology, Dankook University, Cheonan 31116, Republic of Korea')
    affil_run.font.size = Pt(8)
    affil_run.font.name = 'Arial'
    affil_run.italic = True
    affil_para.paragraph_format.space_after = Pt(12)

    # ══════════════════════════════════════════════════════════════════════════
    # ABSTRACT
    # ══════════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, 'Abstract', level=1, font_size=10)

    # Bioinformatics Application Note abstract format:
    # Summary; Availability and Implementation; Contact; Supplementary information.
    summary_para = doc.add_paragraph()
    run = summary_para.add_run('Summary: ')
    run.font.size = Pt(9); run.font.name = 'Arial'; run.bold = True
    run = summary_para.add_run(
        "IlluMeta is an automated dual-pipeline framework (minfi and SeSAMe) for Illumina methylation arrays "
        "that generates a conservative consensus DMP set and adapts batch correction to sample size using CRF."
    )
    run.font.size = Pt(9); run.font.name = 'Arial'

    avail_para = doc.add_paragraph()
    run = avail_para.add_run('Availability and Implementation: ')
    run.font.size = Pt(9); run.font.name = 'Arial'; run.bold = True
    version = get_illumeta_version()
    ver_note = f" (v{version})" if version else ""
    run = avail_para.add_run(
        f"Source code{ver_note} is available at {repo_url} (Apache 2.0). IlluMeta provides a Python CLI that orchestrates "
        "an R-based analysis workflow from raw IDATs to an interactive HTML dashboard."
    )
    run.font.size = Pt(9); run.font.name = 'Arial'

    contact_para = doc.add_paragraph()
    run = contact_para.add_run('Contact: ')
    run.font.size = Pt(9); run.font.name = 'Arial'; run.bold = True
    contact_disp = contact_email.strip() if contact_email else '[ADD_EMAIL]'
    run = contact_para.add_run(contact_disp)
    run.font.size = Pt(9); run.font.name = 'Arial'

    supp_para = doc.add_paragraph()
    run = supp_para.add_run('Supplementary information: ')
    run.font.size = Pt(9); run.font.name = 'Arial'; run.bold = True
    run = supp_para.add_run(supplementary_text)
    run.font.size = Pt(9); run.font.name = 'Arial'
    supp_para.paragraph_format.space_after = Pt(8)

    # ══════════════════════════════════════════════════════════════════════════
    # 1. INTRODUCTION
    # ══════════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, '1  Introduction', level=1, font_size=10)

    add_body_with_refs(doc,
        "Illumina Infinium BeadChip microarrays (450K, EPIC, and EPIC v2) remain the most widely "
        "used platform for genome-wide DNA methylation profiling, with thousands of datasets "
        "deposited in public repositories such as the Gene Expression Omnibus (GEO) [1]. Standard analysis "
        "pipelines, including minfi [2] and ChAMP [3], provide comprehensive workflows for "
        "quality control, normalization, and differential methylation testing. However, each "
        "pipeline applies a different preprocessing strategy\u2014for example, Noob background correction "
        "and/or functional normalization in minfi versus BMIQ in ChAMP\u2014and users typically select one pipeline without a "
        "principled way to assess whether detected signals are robust to methodological choice."
    )

    add_body_with_refs(doc,
        "Batch effects represent another persistent challenge. When sample processing spans "
        "multiple Sentrix slides or dates, systematic technical variation can dwarf biological "
        "signal [4]. Surrogate Variable Analysis (SVA) [5] and ComBat [6] are effective "
        "batch-correction methods, but their applicability depends on sample size: SVA requires "
        "sufficient degrees of freedom, and ComBat can over-fit when groups are small. Existing "
        "tools leave users to make these choices manually, with no adaptive guidance."
    )

    add_body_with_refs(doc,
        "To address these gaps, we developed IlluMeta, a Python/R framework that (i) executes "
        "two independent normalization pipelines in parallel (minfi Noob and SeSAMe pOOBAH [7]), "
        "(ii) intersects DMP results to yield a high-confidence consensus set, and (iii) applies "
        "a Correction Robustness Framework (CRF) that adaptively enables or disables "
        "statistical features based on sample size. IlluMeta automates the full workflow from "
        "IDAT file ingestion to an interactive HTML dashboard, providing reproducible, "
        "publication-ready output with minimal user intervention."
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 2. METHODS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, '2  Methods', level=1, font_size=10)

    # 2.1 Dual-pipeline design
    add_heading_styled(doc, '2.1  Dual-pipeline design', level=2, font_size=9)
    add_body_with_refs(doc,
        "IlluMeta runs two normalization pipelines in parallel. The first uses minfi [2] with "
        "Noob (normal-exponential out-of-band) background correction (preprocessNoob). "
        "The second uses SeSAMe [7] with pOOBAH (p-value with OOB probes for Array Hybridization) "
        "masking and Noob correction. Both pipelines independently perform detection p-value "
        "filtering, cross-reactive probe removal [8], sex chromosome exclusion, and SNP probe "
        "filtering."
    )

    # 2.2 Batch correction and CRF
    add_heading_styled(doc, '2.2  Batch correction and CRF', level=2, font_size=9)
    add_body_with_refs(doc,
        "Each pipeline independently detects batch variables (e.g., Sentrix_ID) and evaluates "
        "confounding with the variable of interest. If batch correction is warranted, the "
        "framework automatically selects between SVA [5], ComBat [6], or limma removeBatchEffect "
        "[9] using a decision tree that prioritizes SVA for datasets with sufficient degrees of "
        "freedom and falls back to ComBat or limma for smaller studies. The Correction Robustness "
        "Framework (CRF) defines four sample-size tiers\u2014minimal (n < 12 or per-group < 6), "
        "small (12\u201323), moderate (24\u201349), and large (n \u2265 50)\u2014and progressively enables "
        "advanced diagnostics (variance partition, multi-method comparison, signal-stability "
        "scoring) as sample size increases (Figure 2D). This prevents over-correction artifacts "
        "in underpowered studies while providing comprehensive quality assurance for larger cohorts. "
        "IlluMeta also reports a Correction Adequacy Framework (CAF) score (0\u20131) that summarizes "
        "calibration, signal preservation, and batch-removal performance."
    )

    # 2.3 Consensus intersection
    add_heading_styled(doc, '2.3  Consensus intersection', level=2, font_size=9)
    add_body_text(doc,
        "After each pipeline identifies significant DMPs (FDR < 0.05 via limma), IlluMeta "
        "intersects the two DMP lists to produce a consensus set. Only CpG sites that reach "
        "significance in both pipelines and show concordant direction of effect are retained. "
        "This dual-confirmation strategy reduces pipeline-specific false positives while "
        "preserving genuinely robust biological signals."
    )

    # 2.4 Additional features
    add_heading_styled(doc, '2.4  Additional features', level=2, font_size=9)
    add_body_with_refs(doc,
        "IlluMeta integrates cell-type deconvolution via EpiDISH [10] and reference-free "
        "estimation (RefFreeEWAS), epigenetic age estimation through methylclock and planet "
        "[11], DMR detection via dmrff [12], and automated GEO dataset download. All results "
        "are compiled into a self-contained interactive HTML dashboard with downloadable tables, "
        "volcano plots, Manhattan plots, and auto-generated methods text."
    )

    # ══════════════════════════════════════════════════════════════════════════
    # TABLE 1: Feature comparison
    # ══════════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, 'Table 1. Feature comparison of IlluMeta, minfi, and ChAMP', level=2, font_size=9)

    features = [
        ['Feature', 'minfi', 'ChAMP', 'IlluMeta'],
        ['Normalization', 'Noob (preprocessNoob); optional funnorm', 'BMIQ', 'Dual: Noob (minfi) + Noob/pOOBAH (SeSAMe)'],
        ['Dual Pipeline', 'No', 'No', 'Yes'],
        ['Consensus DMP Calling', 'No', 'No', 'Yes'],
        ['Auto Batch Correction', 'Manual', 'Manual', 'Auto (SVA/ComBat/limma)'],
        ['Sample Size Adaptation', 'No', 'No', 'CRF'],
        ['Interactive Dashboard', 'No', 'No', 'Yes'],
        ['GEO Integration', 'Manual', 'Manual', 'Built-in download'],
        ['Auto Methods Text', 'No', 'No', 'Yes'],
        ['Cell Deconvolution', 'FlowSorted', 'RefBaseEWAS', 'EpiDISH + RefFreeEWAS'],
        ['Epigenetic Clocks', 'No', 'No', 'methylclock/planet'],
        ['EPIC v2 Support', 'Manual setup', 'No', 'Auto-detected'],
    ]

    table1 = doc.add_table(rows=len(features), cols=4)
    table1.style = 'Table Grid'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(features):
        for j, cell_text in enumerate(row_data):
            cell = table1.cell(i, j)
            cell.text = ''
            para = cell.paragraphs[0]
            run = para.add_run(cell_text)
            run.font.size = Pt(8)
            run.font.name = 'Arial'
            if i == 0:
                run.bold = True
                set_cell_shading(cell, 'D9E2F3')
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT

    # Set column widths
    for row in table1.rows:
        row.cells[0].width = Cm(4.0)
        row.cells[1].width = Cm(3.0)
        row.cells[2].width = Cm(3.0)
        row.cells[3].width = Cm(4.5)

    doc.add_paragraph()  # spacing

    # ══════════════════════════════════════════════════════════════════════════
    # FIGURE 1
    # ══════════════════════════════════════════════════════════════════════════
    fig1_path = os.path.join(figures_dir, 'Figure1_Workflow.png')
    if os.path.exists(fig1_path):
        fig1_para = doc.add_paragraph()
        fig1_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fig1_run = fig1_para.add_run()
        fig1_run.add_picture(fig1_path, width=Inches(5.5))

        cap1 = doc.add_paragraph()
        cap1_run = cap1.add_run(
            'Figure 1. IlluMeta workflow overview. '
            'Input IDAT files or GEO accessions are processed through dual normalization '
            'pipelines (minfi Noob and SeSAMe pOOBAH), with automated batch correction and '
            'consensus DMP intersection. The CRF module adaptively controls statistical '
            'features based on sample size. Output is delivered as an interactive HTML dashboard.'
        )
        cap1_run.font.size = Pt(8)
        cap1_run.font.name = 'Arial'
        cap1_run.italic = True
        cap1.paragraph_format.space_after = Pt(8)
    else:
        add_body_text(doc, '[Figure 1: Workflow diagram - file not found]')

    # ══════════════════════════════════════════════════════════════════════════
    # 3. RESULTS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, '3  Results', level=1, font_size=10)

    # Sort benchmark data by platform then sample size for consistent presentation
    platform_order = {'450k': 0, 'EPIC': 1, 'EPICv2': 2}
    benchmark_data = sorted(benchmark_data,
                            key=lambda r: (platform_order.get(r.get('platform', ''), 9),
                                           safe_int(r.get('n_samples', 0))))

    # Build dynamic results text
    n_datasets = len(benchmark_data)
    total_samples = sum(safe_int(d.get('n_samples', 0)) for d in benchmark_data)
    platforms_list = sorted(set(d.get('platform', '') for d in benchmark_data if d.get('platform')),
                            key=lambda p: platform_order.get(p, 9))
    if len(platforms_list) > 2:
        platform_str = ', '.join(platforms_list[:-1]) + ', and ' + platforms_list[-1]
    else:
        platform_str = ' and '.join(platforms_list)
    platforms_disp = []
    for p in platforms_list:
        if str(p).lower() == "450k":
            platforms_disp.append("450K")
        elif str(p) == "EPICv2":
            platforms_disp.append("EPIC v2")
        else:
            platforms_disp.append(str(p))
    if len(platforms_disp) > 2:
        platform_disp_str = ", ".join(platforms_disp[:-1]) + ", and " + platforms_disp[-1]
    else:
        platform_disp_str = " and ".join(platforms_disp)
    sample_range = sorted([safe_int(d.get('n_samples', 0)) for d in benchmark_data])

    add_body_with_refs(doc,
        f"We benchmarked IlluMeta on {n_datasets} publicly available GEO datasets "
        f"(total n = {total_samples}; range {sample_range[0]}\u2013{sample_range[-1]}) spanning "
        f"{platform_disp_str} arrays and CRF tiers from minimal to large (Table 2). "
        f"The benchmarks cover diverse study designs including breast cancer tissues and cell lines "
        f"(e.g., DCIS or tumor vs normal, drug-resistance models), placenta-related pregnancy outcomes "
        f"(e.g., PTL/PPROM and fetal growth restriction cohorts), and blood-based case-control studies."
    )

    # Describe key findings per dataset
    datasets_with_dmps = [d for d in benchmark_data if safe_int(d.get('consensus_dmps', 0)) > 0]
    datasets_no_dmps = [d for d in benchmark_data if safe_int(d.get('consensus_dmps', 0)) == 0]

    if datasets_with_dmps:
        dmp_text_parts = []
        for d in datasets_with_dmps:
            cdmps = safe_int(d.get('consensus_dmps', 0))
            mdmps = safe_int(d.get('minfi_dmps', 0))
            sdmps = safe_int(d.get('sesame_dmps', 0))
            gse = d.get('gse_id', '')
            n = d.get('n_samples', '?')
            plat_raw = d.get('platform', '')
            if str(plat_raw).lower() == "450k":
                plat = "450K"
            elif str(plat_raw) == "EPICv2":
                plat = "EPIC v2"
            else:
                plat = str(plat_raw)
            if mdmps > 0 and sdmps > 0:
                concordance = cdmps / min(mdmps, sdmps) * 100
                dmp_text_parts.append(
                    f"{gse} ({plat}, n = {n}) yielded {cdmps:,} consensus DMPs "
                    f"({concordance:.0f}% concordance between pipelines)"
                )
        if dmp_text_parts:
            add_body_text(doc,
                'Among datasets with detectable signal, ' +
                '. '.join(dmp_text_parts) +
                '. The high concordance rates indicate that both normalization pipelines '
                'converge on the same biological signal, supporting the dual-pipeline consensus '
                'approach (Figure 2A,C).'
            )

    if datasets_no_dmps:
        no_dmp_parts = []
        for d in datasets_no_dmps:
            plat_raw = d.get('platform', '')
            if str(plat_raw).lower() == "450k":
                plat = "450K"
            elif str(plat_raw) == "EPICv2":
                plat = "EPIC v2"
            else:
                plat = str(plat_raw)
            no_dmp_parts.append(f"{d.get('gse_id', '')} ({plat}, n = {d.get('n_samples', '?')})")
        add_body_text(doc,
            f"In contrast, {', '.join(no_dmp_parts)} yielded zero consensus DMPs after FDR "
            f"correction, reflecting either limited robust signal at the configured thresholds or "
            f"cross-pipeline disagreement. In these cases, IlluMeta's CRF tiers limited batch-correction "
            f"complexity to reduce over-correction risk, and the dashboard still reports full per-pipeline "
            f"results plus diagnostics for sensitivity analyses."
        )

    # CAF discussion
    caf_vals = [safe_float(d.get('caf_score', 0), 0) for d in benchmark_data if d.get('caf_score') is not None]
    caf_vals = [v for v in caf_vals if isinstance(v, (int, float))]
    if caf_vals:
        add_body_text(doc,
            f"IlluMeta summarizes overall correction adequacy using the Correction Adequacy Framework (CAF) "
            f"score (Figure 2B), which integrates calibration, signal preservation, and batch reduction. "
            f"Across the {n_datasets} benchmark datasets, CAF scores ranged from {min(caf_vals):.3f} to "
            f"{max(caf_vals):.3f} (higher = better)."
        )

    # EPIC v2 specific highlight
    epicv2_datasets = [d for d in benchmark_data if d.get('platform') == 'EPICv2']
    if epicv2_datasets:
        d = epicv2_datasets[0]
        add_body_text(doc,
            f"Notably, {d.get('gse_id', '')} demonstrates IlluMeta's seamless support for the "
            f"latest EPIC v2 array (930K probes). The framework automatically detected the array "
            f"type, loaded the appropriate EPICv2 manifest and annotation packages, and produced "
            f"{safe_int(d.get('consensus_dmps', 0)):,} consensus DMPs with high pipeline concordance, "
            f"confirming that IlluMeta is immediately applicable to next-generation array data "
            f"without manual configuration."
        )

    # ══════════════════════════════════════════════════════════════════════════
    # TABLE 2: Benchmark summary
    # ══════════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, f'Table 2. Benchmark results across {n_datasets} GEO datasets', level=2, font_size=9)

    headers2 = ['Dataset', 'Platform', 'n', 'CRF\nTier',
                'Minfi\nDMPs', 'SeSAMe\nDMPs', 'Consensus\nDMPs',
                'CAF']
    n_cols = len(headers2)
    n_rows = 1 + len(benchmark_data)

    table2 = doc.add_table(rows=n_rows, cols=n_cols)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for j, h in enumerate(headers2):
        cell = table2.cell(0, j)
        cell.text = ''
        para = cell.paragraphs[0]
        run = para.add_run(h)
        run.font.size = Pt(6)
        run.font.name = 'Arial'
        run.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'D9E2F3')

    # Data rows
    for i, d in enumerate(benchmark_data):
        plat_raw = d.get('platform', '')
        if str(plat_raw).lower() == "450k":
            plat = "450K"
        elif str(plat_raw) == "EPICv2":
            plat = "EPIC v2"
        else:
            plat = str(plat_raw)
        row_data = [
            d.get('gse_id', ''),
            plat,
            str(d.get('n_samples', '')),
            d.get('crf_tier', '').capitalize(),
            f"{safe_int(d.get('minfi_dmps', 0)):,}",
            f"{safe_int(d.get('sesame_dmps', 0)):,}",
            f"{safe_int(d.get('consensus_dmps', 0)):,}",
            f"{safe_float(d.get('caf_score', 0)):.3f}",
        ]
        for j, val in enumerate(row_data):
            cell = table2.cell(i + 1, j)
            cell.text = ''
            para = cell.paragraphs[0]
            run = para.add_run(val)
            run.font.size = Pt(7)
            run.font.name = 'Arial'
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Table footnote
    footnote_para = doc.add_paragraph()
    fn_run = footnote_para.add_run(
        'DMPs: FDR < 0.05. '
        'Consensus: DMPs significant in both pipelines with concordant direction. '
        'CAF: Correction Adequacy Framework score (higher = better). '
        'See Supplementary Table S1 for nominal DMPs, lambda, and DMR counts.'
    )
    fn_run.font.size = Pt(7)
    fn_run.font.name = 'Arial'
    fn_run.italic = True
    footnote_para.paragraph_format.space_after = Pt(8)

    # ══════════════════════════════════════════════════════════════════════════
    # FIGURE 2
    # ══════════════════════════════════════════════════════════════════════════
    fig2_path = os.path.join(figures_dir, 'Figure2_Benchmark.png')
    if os.path.exists(fig2_path):
        fig2_para = doc.add_paragraph()
        fig2_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fig2_run = fig2_para.add_run()
        fig2_run.add_picture(fig2_path, width=Inches(5.5))

        cap2 = doc.add_paragraph()
        cap2_run = cap2.add_run(
            f'Figure 2. Benchmark summary across {n_datasets} GEO datasets spanning {platform_disp_str} arrays. '
            '(A) DMP counts by pipeline (Minfi, SeSAMe) and consensus intersection, sorted by platform '
            '(capped y-axis; extreme values are annotated). '
            '(B) Correction Adequacy Framework (CAF) score (higher = better). '
            '(C) Consensus DMP counts with pipeline concordance rates (capped y-axis; values annotated). '
            '(D) CRF feature availability by sample-size tier. '
            'Y = enabled, \u2248 = limited, N = disabled.'
        )
        cap2_run.font.size = Pt(8)
        cap2_run.font.name = 'Arial'
        cap2_run.italic = True
        cap2.paragraph_format.space_after = Pt(8)
    else:
        add_body_text(doc, '[Figure 2: Benchmark multi-panel - file not found]')

    # ══════════════════════════════════════════════════════════════════════════
    # 4. DISCUSSION
    # ══════════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, '4  Discussion', level=1, font_size=10)

    add_body_with_refs(doc,
        "IlluMeta addresses two fundamental challenges in Illumina methylation analysis: "
        "pipeline-dependence of DMP calls and sample-size-inappropriate batch correction. "
        "By running minfi and SeSAMe in parallel and intersecting their results, IlluMeta "
        "provides a built-in replication mechanism that filters pipeline-specific artifacts "
        "without requiring external validation cohorts. The CRF tier system ensures "
        "that batch-correction complexity scales with statistical power, preventing "
        "over-correction in small-sample exploratory studies while enabling comprehensive "
        "quality assurance in larger cohorts."
    )

    concordance_vals = []
    for d in datasets_with_dmps:
        cdmps = safe_int(d.get('consensus_dmps', 0))
        mdmps = safe_int(d.get('minfi_dmps', 0))
        sdmps = safe_int(d.get('sesame_dmps', 0))
        if mdmps > 0 and sdmps > 0:
            concordance_vals.append(cdmps / min(mdmps, sdmps) * 100)
    concord_txt = ""
    if concordance_vals:
        concord_txt = f" In signal-bearing datasets, concordance ranged from {min(concordance_vals):.0f}% to {max(concordance_vals):.0f}%."

    add_body_with_refs(doc,
        f"Our benchmark across {n_datasets} datasets spanning {platform_disp_str} arrays and sample sizes "
        f"from {sample_range[0]} to {sample_range[-1]} demonstrates that the consensus approach "
        f"produces conservative, high-confidence DMP sets.{concord_txt} "
        f"In datasets with zero consensus, IlluMeta reports no cross-pipeline hits while still providing "
        f"pipeline-specific results for transparency and sensitivity analyses."
    )

    add_body_text(doc,
        "IlluMeta's automatic platform detection and manifest handling enables seamless analysis of "
        "multiple Illumina array generations without manual configuration (Table 1), including EPIC v2 support. "
        "The Correction Adequacy Framework (CAF) score provides users with a single summary metric of overall "
        "analysis quality, analogous to alignment quality scores in genomics."
    )

    add_body_text(doc,
        "Limitations include the current restriction to two-group comparisons and the reliance "
        "on limma for differential testing. Future versions will support multi-group designs, "
        "longitudinal analyses, and integration with long-read methylation data. Additionally, "
        "the consensus approach is inherently conservative: genuine DMPs detected by only one "
        "pipeline are excluded, which may reduce sensitivity in datasets with subtle effects."
    )

    add_body_text(doc,
        "In summary, IlluMeta provides an automated, reproducible, and statistically principled "
        "framework for DNA methylation analysis that is suitable for both experienced "
        "bioinformaticians and researchers new to epigenomics."
    )

    # ══════════════════════════════════════════════════════════════════════════
    # FUNDING
    # ══════════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, 'Funding', level=1, font_size=10)
    add_body_text(doc,
        "This work was supported by the National Research Foundation of Korea (NRF) grant "
        "funded by the Korea government (MSIT)."
    )

    # ══════════════════════════════════════════════════════════════════════════
    # REFERENCES (numbered, with PMIDs)
    # ══════════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, 'References', level=1, font_size=10)

    references = [
        (
            "Barrett,T. et al. (2013) NCBI GEO: archive for functional genomics data "
            "sets\u2014update. Nucleic Acids Res., 41, D991\u2013D995.",
            "PMID: 23193258"
        ),
        (
            "Aryee,M.J. et al. (2014) Minfi: a flexible and comprehensive Bioconductor "
            "package for the analysis of Infinium DNA methylation microarrays. "
            "Bioinformatics, 30, 1363\u20131369.",
            "PMID: 24478339"
        ),
        (
            "Morris,T.J. et al. (2014) ChAMP: 450k Chip Analysis Methylation Pipeline. "
            "Bioinformatics, 30, 428\u2013430.",
            "PMID: 24336642"
        ),
        (
            "Leek,J.T. et al. (2010) Tackling the widespread and critical impact of batch "
            "effects in high-throughput data. Nat. Rev. Genet., 11, 733\u2013739.",
            "PMID: 20838408"
        ),
        (
            "Leek,J.T. and Storey,J.D. (2007) Capturing heterogeneity in gene expression "
            "studies by surrogate variable analysis. PLoS Genet., 3, e161.",
            "PMID: 17907809"
        ),
        (
            "Johnson,W.E. et al. (2007) Adjusting batch effects in microarray expression "
            "data using empirical Bayes methods. Biostatistics, 8, 118\u2013127.",
            "PMID: 16632515"
        ),
        (
            "Zhou,W. et al. (2018) SeSAMe: reducing artifactual detection of DNA methylation "
            "by Infinium BeadChips in genomic deletions. Nucleic Acids Res., 46, e123.",
            "PMID: 30085201"
        ),
        (
            "Pidsley,R. et al. (2016) Critical evaluation of the Illumina MethylationEPIC "
            "BeadChip microarray for whole-genome DNA methylation profiling. Genome Biol., "
            "17, 208.",
            "PMID: 27717381"
        ),
        (
            "Ritchie,M.E. et al. (2015) limma powers differential expression analyses for "
            "RNA-sequencing and microarray studies. Nucleic Acids Res., 43, e47.",
            "PMID: 25605792"
        ),
        (
            "Teschendorff,A.E. et al. (2017) A comparison of reference-based algorithms "
            "for correcting cell-type heterogeneity in Epigenome-Wide Association Studies. "
            "BMC Bioinformatics, 18, 105.",
            "PMID: 28193155"
        ),
        (
            "Pelegri-Siso,D. et al. (2021) methylclock: a Bioconductor package to estimate "
            "DNA methylation age. Bioinformatics, 37, 1759\u20131760.",
            "PMID: 32960939"
        ),
        (
            "Suderman,M. et al. (2018) dmrff: identifying differentially methylated "
            "regions efficiently with power and control. bioRxiv, 508556.",
            "doi: 10.1101/508556"
        ),
    ]

    for i, (ref_text, pmid) in enumerate(references, 1):
        ref_para = doc.add_paragraph()
        ref_para.paragraph_format.space_after = Pt(1)
        ref_para.paragraph_format.space_before = Pt(0)
        ref_para.paragraph_format.line_spacing = Pt(10)

        # Number
        num_run = ref_para.add_run(f'[{i}] ')
        num_run.font.size = Pt(7.5)
        num_run.font.name = 'Arial'
        num_run.bold = True

        # Reference text
        text_run = ref_para.add_run(ref_text + ' ')
        text_run.font.size = Pt(7.5)
        text_run.font.name = 'Arial'

        # PMID
        pmid_run = ref_para.add_run(pmid)
        pmid_run.font.size = Pt(7.5)
        pmid_run.font.name = 'Arial'
        pmid_run.bold = True
        pmid_run.font.color.rgb = RGBColor(0, 0, 180)

    # ══════════════════════════════════════════════════════════════════════════
    # SAVE
    # ══════════════════════════════════════════════════════════════════════════
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    print(f"\nSaved Application Note: {out_path}")
    file_size = os.path.getsize(out_path)
    print(f"File size: {file_size:,} bytes")


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description='Build Application Note DOCX')
    parser.add_argument('--results-dir', default='benchmarks/application_note_submission_results',
                        help='Directory containing IlluMeta results')
    parser.add_argument('--figures-dir', default='benchmarks/paper_figures',
                        help='Directory containing generated figures')
    parser.add_argument('--out', default='benchmarks/application_note/IlluMeta_Application_Note.docx',
                        help='Output DOCX path')
    parser.add_argument('--contact-email', default='',
                        help='Contact email to include in the structured abstract (required for submission).')
    parser.add_argument('--repo-url', default='https://github.com/kangk1204/illumeta',
                        help='Repository URL for Availability and Implementation.')
    parser.add_argument('--supplementary-text',
                        default='Supplementary data are available at Bioinformatics online.',
                        help='Text for the abstract Supplementary information field.')
    parser.add_argument('--allow-fallback', action='store_true',
                        help='Allow building a DOCX with hardcoded placeholder benchmark values when results are missing (NOT for submission).')
    args = parser.parse_args()

    print("IlluMeta Application Note DOCX Builder")
    print("=" * 50)

    # Collect benchmark data
    print("\nCollecting benchmark data...")
    benchmark_data = collect_benchmark_data(args.results_dir)
    print(f"Found {len(benchmark_data)} datasets")

    if not benchmark_data:
        if not args.allow_fallback:
            sys.exit(
                "ERROR: No benchmark data found under --results-dir. "
                "Provide a valid results directory (with GSE*/.../summary.json), "
                "or pass --allow-fallback to generate a non-submission placeholder DOCX."
            )
        print("WARNING: No benchmark data found. Building DOCX with placeholder benchmark values (--allow-fallback).")
        benchmark_data = [
            {'gse_id': 'GSE113687', 'platform': 'EPIC', 'n_samples': '20', 'n_con': 5, 'n_test': 15,
             'crf_tier': 'minimal', 'minfi_dmps': 205948, 'sesame_dmps': 193633,
             'consensus_dmps': 171391, 'minfi_dmrs': 14537, 'sesame_dmrs': 13979,
             'lambda_minfi': 2.013, 'lambda_sesame': 1.145,
             'lambda_minfi_raw': 7.215, 'caf_score': 0.6001},
            {'gse_id': 'GSE66313', 'platform': '450k', 'n_samples': '29', 'n_con': 14, 'n_test': 15,
             'crf_tier': 'moderate', 'minfi_dmps': 0, 'sesame_dmps': 0,
             'consensus_dmps': 0, 'minfi_dmrs': 2321, 'sesame_dmrs': 1741,
             'lambda_minfi': 1.063, 'lambda_sesame': 1.099,
             'caf_score': 0.5051},
            {'gse_id': 'GSE74071', 'platform': '450k', 'n_samples': '6', 'n_con': 3, 'n_test': 3,
             'crf_tier': 'minimal', 'minfi_dmps': 0, 'sesame_dmps': 0,
             'consensus_dmps': 0, 'minfi_dmrs': 3956, 'sesame_dmrs': 1234,
             'lambda_minfi': 6.598, 'lambda_sesame': 4.929,
             'caf_score': 0.4670},
            {'gse_id': 'GSE308134', 'platform': 'EPIC', 'n_samples': '25', 'n_con': 12, 'n_test': 13,
             'crf_tier': 'moderate', 'minfi_dmps': 1400, 'sesame_dmps': 1493,
             'consensus_dmps': 807, 'minfi_dmrs': 2394, 'sesame_dmrs': 3384,
             'lambda_minfi': 0.463, 'lambda_sesame': 1.601,
             'caf_score': 0.5957},
            {'gse_id': 'GSE281691', 'platform': 'EPIC', 'n_samples': '160', 'n_con': 80, 'n_test': 80,
             'crf_tier': 'large', 'minfi_dmps': 0, 'sesame_dmps': 0,
             'consensus_dmps': 0, 'minfi_dmrs': 2205, 'sesame_dmrs': 874,
             'lambda_minfi': 0.959, 'lambda_sesame': 0.900,
             'caf_score': 0.4673},
            {'gse_id': 'GSE312914', 'platform': 'EPICv2', 'n_samples': '147', 'n_con': 42, 'n_test': 105,
             'crf_tier': 'large', 'minfi_dmps': 312, 'sesame_dmps': 314,
             'consensus_dmps': 306, 'minfi_dmrs': 4486, 'sesame_dmrs': 4453,
             'lambda_minfi': 0.952, 'lambda_sesame': 0.888,
             'lambda_minfi_raw': 1.810, 'caf_score': 0.4335},
        ]

    # Build DOCX
    print("\nBuilding DOCX...")
    if not args.contact_email:
        print("WARNING: --contact-email not provided; inserting placeholder [ADD_EMAIL].")
    build_docx(
        benchmark_data,
        args.figures_dir,
        args.out,
        contact_email=args.contact_email,
        repo_url=args.repo_url,
        supplementary_text=args.supplementary_text,
    )
    print("\nDone!")


if __name__ == '__main__':
    main()
