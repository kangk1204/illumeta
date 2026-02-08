#!/usr/bin/env python3
"""
Build Supplementary Data DOCX for IlluMeta Application Note.
Generates a comprehensive ~20-30 page A4 supplement covering:
  S1. Introduction & Overview
  S2. Getting Started
  S3. Analysis Pipeline (Detailed Methods)
  S4. Dashboard Navigation Guide
  S5. Output File Reference
  S6. Interpreting Results (Beginner Guide)
  S7. Benchmark Results
  S8. Command-Line Reference
  S9. FAQ & Troubleshooting
  S10. References
"""
import argparse
import csv
import json
import os
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    sys.exit(
        "Missing python-docx.\n"
        "Install (inside your conda env): python -m pip install -r requirements-paper.txt\n"
        "Or (conda-forge): conda install -c conda-forge python-docx"
    )


# ── Helper utilities ──────────────────────────────────────────────────────────

def find_results_dir(gse_dir):
    """Find the results directory inside a GSE folder (handles varied naming)."""
    gse_path = Path(gse_dir)
    std = gse_path / 'illumeta_results'
    if std.exists() and (std / 'summary.json').exists():
        return std
    for d in sorted(gse_path.iterdir()):
        if d.is_dir() and d.name.endswith('_results') and (d / 'summary.json').exists():
            return d
    return None


def safe_float(val, default=0.0):
    try:
        if val is None or val == '' or val == 'NA':
            return default
        return float(val)
    except (ValueError, TypeError):
        return default


def safe_int(val, default=0):
    try:
        if val is None or val == '' or val == 'NA':
            return default
        return int(float(val))
    except (ValueError, TypeError):
        return default


def count_fdr_significant_csv_rows(csv_path, threshold=0.05):
    """Count rows with FDR/q-value < threshold for common column names."""
    fdr_cols = ['p.adjust', 'FDR', 'adj.P.Val', 'qvalue', 'q.value']
    try:
        with open(csv_path, 'r') as f:
            reader = csv.DictReader(f)
            if reader.fieldnames is None:
                return 0
            cols = set(reader.fieldnames)
            chosen = next((c for c in fdr_cols if c in cols), None)
            if chosen is None:
                return 0
            sig = 0
            for r in reader:
                try:
                    v = r.get(chosen, '')
                    if v is None or v == '' or v == 'NA':
                        continue
                    if float(v) < threshold:
                        sig += 1
                except (ValueError, TypeError):
                    continue
            return sig
    except OSError:
        return 0


def find_showcase_dataset(results_dir, preferred_gse=None):
    """Find a showcase dataset for embedding example figures.

    If preferred_gse is given and exists, use it. Otherwise fall back to the
    dataset with the highest CAF score among those with consensus DMPs > 0.

    Returns the results subdirectory (containing PNGs, summary.json, etc.) or None.
    """
    results_path = Path(results_dir)
    if not results_path.exists():
        return None

    # If a preferred GSE is specified, try it first
    if preferred_gse:
        pref_dir = results_path / preferred_gse
        if pref_dir.is_dir():
            res_dir = find_results_dir(pref_dir)
            if res_dir and (res_dir / 'summary.json').exists():
                print(f"Showcase dataset: {preferred_gse} (user-specified)")
                return res_dir
        print(f"WARNING: Preferred showcase {preferred_gse} not found, falling back to auto-selection.")

    # Auto-select: highest CAF score among datasets with consensus DMPs > 0
    best_dir = None
    best_caf = -1
    best_name = ''

    for gse_dir in sorted(results_path.glob('GSE*')):
        if not gse_dir.is_dir():
            continue
        res_dir = find_results_dir(gse_dir)
        if res_dir is None:
            continue

        sj_file = res_dir / 'summary.json'
        if not sj_file.exists():
            continue

        with open(sj_file, 'r') as f:
            sj = json.load(f)
        consensus = sj.get('intersect_up', 0) + sj.get('intersect_down', 0)
        caf = safe_float(sj.get('primary_caf_score', 0))

        if consensus > 0 and caf > best_caf:
            best_caf = caf
            best_dir = res_dir
            best_name = gse_dir.name

    if best_dir:
        print(f"Showcase dataset: {best_name} (CAF={best_caf:.3f}, auto-selected)")
    return best_dir


def collect_benchmark_data(results_dir):
    """Collect benchmark metrics from results directories."""
    summary = []
    results_path = Path(results_dir)
    if not results_path.exists():
        return summary

    for gse_dir in sorted(results_path.glob('GSE*')):
        if not gse_dir.is_dir():
            continue
        res_dir = find_results_dir(gse_dir)
        if res_dir is None:
            continue

        gse_id = gse_dir.name
        row = {'gse_id': gse_id}

        # summary.json
        sj_file = res_dir / 'summary.json'
        if sj_file.exists():
            with open(sj_file, 'r') as f:
                sj = json.load(f)
                row['n_con'] = sj.get('n_con', 0)
                row['n_test'] = sj.get('n_test', 0)
                row['caf_score'] = sj.get('primary_caf_score', 0)
                row['crf_tier'] = sj.get('crf_sample_tier', '')
                n_total = sj.get('n_con', 0) + sj.get('n_test', 0)
                row['n_samples'] = str(n_total)
                row['consensus_dmps'] = (sj.get('intersect_up', 0) + sj.get('intersect_down', 0))

        # CRF_Sample_Tier.csv (tier only; n_samples uses post-QC count from summary.json)
        tier_file = res_dir / 'CRF_Sample_Tier.csv'
        if tier_file.exists():
            with open(tier_file, 'r') as f:
                reader = csv.DictReader(f)
                for r in reader:
                    row['crf_tier'] = r.get('tier', row.get('crf_tier', ''))
                    break

        # Platform from Probe_Filter_Summary.csv
        probe_file = res_dir / 'Probe_Filter_Summary.csv'
        if probe_file.exists():
            with open(probe_file, 'r') as f:
                reader = csv.DictReader(f)
                for r in reader:
                    if r.get('step') == 'raw':
                        n_probes = safe_int(r.get('remaining', 0))
                        if n_probes > 900000:
                            row['platform'] = 'EPICv2'
                        elif n_probes > 600000:
                            row['platform'] = 'EPIC'
                        else:
                            row['platform'] = '450k'
                        break

        # Lambda from ablation summary (corrected value)
        for pipeline in ['Minfi', 'Sesame']:
            abl_file = res_dir / f'{pipeline}_Ablation_Summary.csv'
            if abl_file.exists():
                with open(abl_file, 'r') as f:
                    reader = csv.DictReader(f)
                    for r in reader:
                        if r.get('metric') == 'lambda':
                            row[f'lambda_{pipeline.lower()}'] = safe_float(
                                r.get('corrected', r.get('raw', '1.0')), 1.0)
                            break

        # DMP counts (both FDR<0.05 and nominal: raw P<0.05 + |logFC|>1)
        for pipeline in ['Minfi', 'Sesame']:
            dmp_file = res_dir / f'{pipeline}_DMPs_full.csv'
            if dmp_file.exists():
                fdr_count = 0
                nominal_count = 0
                with open(dmp_file, 'r') as f:
                    reader = csv.DictReader(f)
                    for r in reader:
                        try:
                            fdr = float(r.get('adj.P.Val', r.get('FDR', '1')))
                            if fdr < 0.05:
                                fdr_count += 1
                            pval = float(r.get('P.Value', r.get('pval', '1')))
                            logfc = abs(float(r.get('logFC', '0')))
                            if pval < 0.05 and logfc > 1:
                                nominal_count += 1
                        except (ValueError, TypeError):
                            pass
                row[f'{pipeline.lower()}_dmps'] = fdr_count
                row[f'{pipeline.lower()}_dmps_nominal'] = nominal_count

        # Consensus DMPs
        consensus_file = res_dir / 'Intersection_Consensus_DMPs.csv'
        if consensus_file.exists():
            with open(consensus_file, 'r') as f:
                row['consensus_dmps'] = sum(1 for _ in f) - 1

        # Count DMRs per pipeline
        for pipeline in ['Minfi', 'Sesame']:
            dmr_file = res_dir / f'{pipeline}_DMRs.csv'
            if dmr_file.exists():
                # DMR output includes non-significant regions; count only FDR<0.05.
                row[f'{pipeline.lower()}_dmrs'] = count_fdr_significant_csv_rows(dmr_file, threshold=0.05)

        if row.get('n_samples') and row.get('n_samples') != '0':
            summary.append(row)

    return summary


# ── Style helpers ─────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set table cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex
    })
    shading.append(shading_elem)


def add_heading(doc, text, level=1, font_size=10):
    """Add a section heading."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Arial'
    run.bold = True
    if level == 1:
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(4)
    elif level == 2:
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(3)
    else:
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(2)
    return para


def add_body(doc, text, font_size=9, italic=False):
    """Add body text paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Arial'
    run.italic = italic
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.line_spacing = Pt(12)
    return para


def add_bullet(doc, text, font_size=9, bold_prefix=None):
    """Add a bullet point."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(1)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.line_spacing = Pt(12)
    para.paragraph_format.left_indent = Cm(0.8)
    if bold_prefix:
        run_b = para.add_run(f"\u2022 {bold_prefix}: ")
        run_b.font.size = Pt(font_size)
        run_b.font.name = 'Arial'
        run_b.bold = True
        run_t = para.add_run(text)
        run_t.font.size = Pt(font_size)
        run_t.font.name = 'Arial'
    else:
        run = para.add_run(f"\u2022 {text}")
        run.font.size = Pt(font_size)
        run.font.name = 'Arial'
    return para


def add_code_block(doc, text, font_size=8):
    """Add a code block (monospace, gray background)."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Courier New'
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.left_indent = Cm(0.5)
    return para


def add_table(doc, headers, rows, col_widths=None, header_color='D9E2F3', font_size=8):
    """Add a formatted table."""
    n_cols = len(headers)
    n_rows = 1 + len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = ''
        para = cell.paragraphs[0]
        run = para.add_run(h)
        run.font.size = Pt(font_size)
        run.font.name = 'Arial'
        run.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    # Data rows
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.cell(i + 1, j)
            cell.text = ''
            para = cell.paragraphs[0]
            run = para.add_run(str(val))
            run.font.size = Pt(font_size)
            run.font.name = 'Arial'
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT

    if col_widths:
        for row_obj in table.rows:
            for j, w in enumerate(col_widths):
                if j < len(row_obj.cells):
                    row_obj.cells[j].width = Cm(w)

    doc.add_paragraph()  # spacing
    return table


def add_figure(doc, fig_path, caption, width=5.0):
    """Add a figure with caption."""
    if os.path.exists(fig_path):
        fig_para = doc.add_paragraph()
        fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fig_run = fig_para.add_run()
        fig_run.add_picture(fig_path, width=Inches(width))

        cap_para = doc.add_paragraph()
        cap_run = cap_para.add_run(caption)
        cap_run.font.size = Pt(8)
        cap_run.font.name = 'Arial'
        cap_run.italic = True
        cap_para.paragraph_format.space_after = Pt(6)
        return True
    else:
        add_body(doc, f'[Figure not found: {os.path.basename(fig_path)}]')
        return False


# ── Document sections ─────────────────────────────────────────────────────────

def build_s1(doc, figures_dir):
    """S1. Introduction & Overview"""
    add_heading(doc, 'S1. Introduction & Overview', level=1, font_size=11)

    add_heading(doc, 'S1.1 What is IlluMeta?', level=2, font_size=9)
    add_body(doc,
        "IlluMeta is an open-source Python/R framework for automated, reproducible DNA methylation "
        "analysis from Illumina Infinium BeadChip microarrays (450K, EPIC, EPIC v2). It is designed "
        "for both experienced bioinformaticians and researchers new to epigenomics.")

    add_body(doc,
        "The key innovation of IlluMeta is its dual-pipeline design: it runs two independent "
        "normalization pipelines (minfi Noob and SeSAMe pOOBAH) in parallel, intersects their "
        "results to produce high-confidence consensus DMPs, and applies a sample-size-adaptive "
        "Correction Robustness Framework (CRF) to gate batch-correction features based on "
        "statistical power.")

    add_heading(doc, 'S1.2 Who is IlluMeta for?', level=2, font_size=9)
    add_bullet(doc, "Biologists with methylation array data who want publication-ready results without extensive R coding.",
               bold_prefix="Wet-lab researchers")
    add_bullet(doc, "Students learning EWAS who benefit from step-by-step workflow guidance and interactive dashboards.",
               bold_prefix="Graduate students")
    add_bullet(doc, "Analysts who want time savings from automated dual-pipeline analysis with built-in QC.",
               bold_prefix="Bioinformaticians")
    add_bullet(doc, "Researchers who need reproducible figures, auto-generated methods text, and complete decision logs.",
               bold_prefix="Paper authors")

    add_heading(doc, 'S1.3 System requirements', level=2, font_size=9)
    add_bullet(doc, "Linux (Ubuntu 20.04+), macOS (Apple Silicon or Intel), or Windows 11 (via WSL2)")
    add_bullet(doc, "R 4.3+ (R 4.4+ recommended for EPIC v2 support)")
    add_bullet(doc, "Python 3.8+")
    add_bullet(doc, "Conda/Mamba (recommended) or venv + system R")
    add_bullet(doc, "At least 8 GB RAM (16 GB recommended for large datasets)")
    add_bullet(doc, "pandoc (required for self-contained HTML reports)")

    add_heading(doc, 'S1.4 Workflow overview', level=2, font_size=9)
    fig1 = os.path.join(figures_dir, 'Figure1_Workflow.png')
    add_figure(doc, fig1,
        "Figure S1. IlluMeta workflow overview. Input IDAT files or GEO accessions are processed "
        "through dual normalization pipelines (minfi Noob and SeSAMe pOOBAH), with automated "
        "batch correction and consensus DMP intersection.", width=5.5)


def build_s2(doc):
    """S2. Getting Started"""
    add_heading(doc, 'S2. Getting Started', level=1, font_size=11)

    add_heading(doc, 'S2.1 Installation', level=2, font_size=9)
    add_body(doc, "The recommended installation uses conda to manage all dependencies:")
    add_code_block(doc,
        "git clone https://github.com/kangk1204/illumeta.git\n"
        "cd illumeta\n"
        "conda env create -f environment.yml\n"
        "conda activate illumeta\n"
        "Rscript r_scripts/setup_env.R\n"
        "python3 illumeta.py doctor")
    add_body(doc,
        "The doctor command verifies that all R packages are installed correctly. "
        "For one-line installation, run: ./scripts/install_full.sh")

    add_heading(doc, 'S2.2 Downloading data from GEO', level=2, font_size=9)
    add_body(doc,
        "IlluMeta can automatically download IDAT files from the Gene Expression Omnibus (GEO):")
    add_code_block(doc,
        "python3 illumeta.py download GSE121633 -o projects/GSE121633")
    add_body(doc,
        "This downloads the raw IDAT files and generates a configure.tsv file with sample metadata. "
        "If a GEO series has multiple platforms, specify one with --platform GPL21145.")

    add_heading(doc, 'S2.3 Configure file format', level=2, font_size=9)
    add_body(doc,
        "The configure.tsv file is a tab-delimited file with at least two columns:")
    add_table(doc,
        ['Column', 'Required', 'Description'],
        [
            ['Basename', 'Yes', 'Path to IDAT file pair (without _Grn/_Red suffix)'],
            ['primary_group', 'Yes*', 'Group label (e.g., Control, Case)'],
            ['SampleID', 'No', 'Optional sample identifier'],
            ['sex', 'No', 'Biological sex (M/F) for sex check'],
            ['age', 'No', 'Chronological age for clock comparison'],
            ['tissue', 'No', 'Tissue type for reference-based deconvolution'],
        ],
        col_widths=[3.0, 1.5, 10.0])
    add_body(doc,
        "*primary_group can be auto-populated using --auto-group with --group-column or --group-key.")

    add_heading(doc, 'S2.4 Running analysis', level=2, font_size=9)
    add_code_block(doc,
        "python3 illumeta.py analysis \\\n"
        "  -i projects/GSE121633 \\\n"
        "  --group_con Control --group_test Case")
    add_body(doc,
        "Results are written to the output directory (default: [input]/[test]_vs_[con]_results/) "
        "and an interactive HTML dashboard is generated in the parent directory.")

    add_heading(doc, 'S2.5 Opening the dashboard', level=2, font_size=9)
    add_body(doc,
        "Open the generated HTML file in any modern web browser:")
    add_code_block(doc,
        "# The dashboard file is named: [test]_vs_[con]_results_index.html\n"
        "# Example:\n"
        "open projects/GSE121633/Case_vs_Control_results_index.html")


def build_s3(doc):
    """S3. Analysis Pipeline (Detailed Methods)"""
    add_heading(doc, 'S3. Analysis Pipeline (Detailed Methods)', level=1, font_size=11)

    # S3.1
    add_heading(doc, 'S3.1 Quality Control', level=2, font_size=9)
    add_body(doc,
        "IlluMeta performs multi-step quality control before differential methylation analysis:")
    add_bullet(doc, "Detection p-value filtering operates in two stages: (1) samples where >20% of probes exceed the "
               "detection p-value threshold (default: p > 0.05) are removed; (2) probes that fail the threshold in any "
               "remaining sample are excluded.",
               bold_prefix="Detection P")
    add_bullet(doc, "Known cross-reactive probes are removed using the curated list from Pidsley et al. (2016) and McCartney et al. (2016).",
               bold_prefix="Cross-reactive probes")
    add_bullet(doc, "Probes overlapping known SNPs (MAF > 0.01) are removed to avoid genotype-driven methylation differences.",
               bold_prefix="SNP filtering")
    add_bullet(doc, "X/Y chromosome probes are excluded by default to prevent sex-driven signals.",
               bold_prefix="Sex chromosome exclusion")
    add_bullet(doc, "Predicted sex (from X/Y methylation) is compared with reported sex; mismatched samples are flagged.",
               bold_prefix="Sex check")

    # S3.2
    add_heading(doc, 'S3.2 Normalization', level=2, font_size=9)
    add_body(doc,
        "IlluMeta runs two independent normalization pipelines in parallel:")
    add_bullet(doc, "Noob (normal-exponential out-of-band) background correction via preprocessNoob. "
               "Functional normalization is available in minfi, but IlluMeta defaults to Noob for consistency.",
               bold_prefix="Minfi Noob")
    add_bullet(doc, "Design masking (qualityMask), Infinium-I channel inference, dye-bias correction (dyeBiasL; "
               "optional TypeINorm), pOOBAH masking, and Noob background correction. Masked probes are carried as NA "
               "during beta-value extraction.",
               bold_prefix="SeSAMe + pOOBAH")
    add_body(doc,
        "Both pipelines produce independent beta-value matrices that are carried forward through "
        "all downstream analyses. IlluMeta reports SeSAMe in two complementary views: "
        "(i) a strict complete-case, Minfi-aligned view for conservative cross-validation, and "
        "(ii) a native pOOBAH-masked view that retains probes with limited missingness and imputes "
        "remaining NAs for downstream modeling. We recommend comparing strict vs native as a "
        "sensitivity analysis.")

    # S3.3
    add_heading(doc, 'S3.3 Batch Detection & Correction', level=2, font_size=9)
    add_body(doc,
        "IlluMeta automatically detects batch variables (e.g., Sentrix_ID, Sentrix_Position) and "
        "evaluates their confounding with the variable of interest using chi-squared tests and "
        "Cramer's V. If batch correction is warranted, the framework selects between three methods:")
    add_bullet(doc, "Surrogate Variable Analysis. Discovers hidden confounders by decomposing residual variation. "
               "Best for larger datasets with sufficient degrees of freedom.",
               bold_prefix="SVA")
    add_bullet(doc, "Empirical Bayes batch correction. Effective when batch structure is known and groups are balanced.",
               bold_prefix="ComBat")
    add_bullet(doc, "removeBatchEffect. Lightweight approach for smaller studies or when SVA/ComBat are inappropriate.",
               bold_prefix="limma")
    add_body(doc,
        "The selection is automatic: SVA is preferred when sample size permits (CRF moderate or large "
        "tier), with ComBat as fallback and limma as the safest option for small samples.")

    # S3.4
    add_heading(doc, 'S3.4 Covariate Detection', level=2, font_size=9)
    add_body(doc,
        "IlluMeta performs automatic PCA-based covariate discovery. For each candidate covariate "
        "(metadata columns, cell-type proportions, clock estimates), the framework tests association "
        "with the top principal components of the methylation data. Covariates that are (a) strongly "
        "associated with PCs but (b) not confounded with the primary group variable are included "
        "in the limma model. Covariates that are too strongly associated with the group variable "
        "(potential colliders) are excluded to avoid masking true biological signal.")
    add_body(doc,
        "Cell-type proportions estimated by EpiDISH or RefFreeEWAS (Section S3.8) are automatically "
        "added to the candidate covariate list. Each cell-type proportion undergoes the same governance "
        "checks as metadata covariates: (a) variance check (single-level or near-constant proportions "
        "are dropped), (b) confounding check (Eta-squared test against primary_group; high-confounding "
        "cell types are excluded), and (c) collinearity filtering (|r| >= 0.8 with another covariate "
        "triggers removal of the less informative variable). Cell types passing all checks are added "
        "to the limma design matrix alongside other selected covariates. An overcorrection guard caps "
        "the total number of covariates to preserve degrees of freedom (see Section S3.14).")

    # S3.5
    add_heading(doc, 'S3.5 DMP Analysis', level=2, font_size=9)
    add_body(doc,
        "Differentially methylated positions (DMPs) are identified using limma (Ritchie et al., 2015). "
        "For each pipeline, a linear model is fit on M-values (logit-transformed beta values) with "
        "group as the primary variable and selected covariates/surrogate variables. Moderated t-statistics "
        "are computed via empirical Bayes shrinkage. Default thresholds: FDR < 0.05, |logFC| > 0.5 "
        "(adjustable via --pval, --lfc, --delta-beta flags).")

    # S3.6
    add_heading(doc, 'S3.6 DMR Analysis', level=2, font_size=9)
    add_body(doc,
        "Differentially methylated regions (DMRs) are detected using dmrff (Suderman et al., 2018), "
        "which identifies genomic regions where multiple adjacent CpGs show coordinated differential "
        "methylation. Parameters: maximum gap between CpGs (default 500 bp), p-value cutoff "
        "(default 0.05). If Tier3 confounding is detected, DMRs are emitted from stratified/meta-analysis.")

    # S3.7
    add_heading(doc, 'S3.7 Consensus Intersection', level=2, font_size=9)
    add_body(doc,
        "The consensus intersection identifies CpG sites that reach significance (FDR < 0.05) in "
        "both the minfi and SeSAMe pipelines with concordant direction of effect. This dual-confirmation "
        "strategy reduces pipeline-specific false positives. IlluMeta produces two intersection views:")
    add_bullet(doc, "Uses the Minfi-aligned SeSAMe results (same probe set as Minfi).",
               bold_prefix="Strict intersection")
    add_bullet(doc, "Uses SeSAMe's native pOOBAH-preserved results, which may have fewer probes but stricter quality.",
               bold_prefix="Native intersection")
    add_body(doc,
        "Fisher's combined p-value is computed for each consensus CpG, along with the mean logFC and "
        "mean delta-beta across pipelines.")

    # S3.8
    add_heading(doc, 'S3.8 Cell Composition Deconvolution', level=2, font_size=9)
    add_body(doc,
        "IlluMeta estimates cell-type composition to adjust for cellular heterogeneity:")
    add_bullet(doc, "Reference-based deconvolution using constrained projection. Supports Blood (FlowSorted.Blood.EPIC), "
               "CordBlood, Placenta (planet), and DLPFC tissue references.",
               bold_prefix="EpiDISH")
    add_bullet(doc, "Reference-free estimation using non-negative matrix factorization. Produces latent cell-type "
               "proportions (Cell_Latent1-5) without requiring a tissue-specific reference.",
               bold_prefix="RefFreeEWAS")
    add_bullet(doc, "Placenta-specific cell composition estimation using the planet R package.",
               bold_prefix="planet")
    add_body(doc,
        "Cell-type estimates are automatically included as candidate covariates in the differential "
        "methylation model. IlluMeta checks each cell-type proportion for (a) sufficient variance "
        "(>1 unique value), (b) low confounding with the primary group variable (Eta-squared < "
        "threshold), and (c) collinearity with other covariates (|r| < 0.8). Cell types passing all "
        "checks are added to the limma design matrix alongside other selected covariates. This "
        "automatic mechanism ensures that cellular heterogeneity is accounted for without requiring "
        "the user to manually specify cell-type adjustments.")

    # S3.9
    add_heading(doc, 'S3.9 Epigenetic Clocks', level=2, font_size=9)
    add_body(doc,
        "IlluMeta estimates DNA methylation age using multiple epigenetic clocks via the methylclock "
        "R package: Horvath, Hannum, Levine (PhenoAge), skinHorvath, PedBE, Wu, and telomere length (TL). "
        "For placental samples, gestational age is estimated using the planet package (RPC and CPC clocks). "
        "Clock estimates can optionally be included as candidate covariates in the DMP model (--include-clock-covariates).")

    # S3.10
    add_heading(doc, 'S3.10 CRF Robustness Framework', level=2, font_size=9)
    add_body(doc,
        "The Correction Robustness Framework (CRF) defines four sample-size tiers and "
        "progressively enables statistical features:")
    add_table(doc,
        ['Tier', 'Total n', 'Per-group min', 'SVA', 'PVCA', 'MMC', 'SSS'],
        [
            ['Minimal', '< 12', '3', 'Disabled', 'Disabled', 'Limited', 'Limited'],
            ['Small', '12-23', '6', 'Disabled', 'Enabled', 'Limited', 'Limited'],
            ['Moderate', '24-49', '12', 'Enabled', 'Enabled', 'Enabled', 'Enabled'],
            ['Large', '>= 50', '25', 'Enabled', 'Enabled', 'Enabled', 'Enabled'],
        ],
        col_widths=[2.0, 1.5, 2.0, 2.0, 2.0, 2.0, 2.0])
    add_body(doc,
        "MMC = Multi-Method Comparison (evaluates multiple correction strategies), "
        "SSS = Signal Stability Score (leave-one-out stability assessment).")

    # S3.11
    add_heading(doc, 'S3.11 CAF Score', level=2, font_size=9)
    add_body(doc,
        "The Correction Adequacy Framework (CAF) provides a single summary metric (0\u20131) of overall "
        "analysis quality, computed as a weighted composite of three components:")
    add_body(doc,
        "CAI = 0.40 \u00d7 Calibration + 0.35 \u00d7 Preservation + 0.25 \u00d7 Batch Removal")
    add_bullet(doc,
        "Measures how well the null distribution is calibrated. Computed as: "
        "1 \u2212 |null_FPR \u2212 0.05| / 0.05, where null_FPR is the permutation-based false positive "
        "rate (Section S3.13). A perfectly calibrated analysis yields FPR = 5% and a score of 1.0.",
        bold_prefix="Calibration (weight 0.40)")
    add_bullet(doc,
        "Measures how much biological signal is retained after batch correction. Computed as: "
        "marker_retention \u00d7 logFC_correlation (if a marker list is provided) or just logFC_correlation "
        "(otherwise). logFC_correlation is the Pearson correlation of per-CpG log-fold changes before "
        "and after correction.",
        bold_prefix="Signal Preservation (weight 0.35)")
    add_bullet(doc,
        "Measures how effectively batch-associated variation is removed. Computed as: "
        "(sig_before \u2212 sig_after) / sig_before, where sig_before and sig_after are the number of "
        "significant batch\u2013PC associations before and after correction.",
        bold_prefix="Batch Removal (weight 0.25)")
    add_body(doc,
        "All three components are clamped to [0, 1]. If a component cannot be computed (e.g., no "
        "batch variable detected), it is treated as NA and the remaining weights are renormalized. "
        "The final CAI score is graded on the following scale:")
    add_table(doc,
        ['Grade', 'CAI range', 'Interpretation'],
        [
            ['EXCELLENT', '\u2265 0.85', 'Well-calibrated, signal preserved, batch removed'],
            ['GOOD', '0.70\u20130.84', 'Adequate correction; minor issues possible'],
            ['FAIR', '0.50\u20130.69', 'Partial correction; review batch diagnostics'],
            ['POOR', '< 0.50', 'Correction may be inadequate; investigate manually'],
        ],
        col_widths=[2.5, 2.0, 10.0])

    # S3.12
    add_heading(doc, 'S3.12 Tier3 Meta-Analysis', level=2, font_size=9)
    add_body(doc,
        "When batch variables are fully confounded with the group variable (i.e., all samples in one "
        "group come from one batch), IlluMeta automatically detects this and performs a stratified "
        "analysis within each batch stratum, followed by meta-analysis (fixed or random effects) to "
        "combine results across strata. This Tier3 approach avoids the impossible task of removing "
        "batch effects that are perfectly collinear with the biological variable.")

    # S3.13
    add_heading(doc, 'S3.13 Permutation Testing', level=2, font_size=9)
    add_body(doc,
        "IlluMeta runs N label permutations (default N = 20) to estimate the empirical null false "
        "positive rate (FPR). In each permutation, the primary_group labels are randomly shuffled "
        "while preserving all other metadata, and the full DMP analysis is re-run on the permuted data.")
    add_bullet(doc,
        "The fraction of CpGs reaching significance (FDR < 0.05) under each null permutation is "
        "recorded. The mean null FPR across permutations estimates how often the pipeline produces "
        "false positives by chance alone.",
        bold_prefix="Null FPR")
    add_bullet(doc,
        "The median genomic inflation factor (\u03bb) across null permutations provides a reference "
        "for the expected lambda under the null hypothesis.",
        bold_prefix="Null lambda")
    add_bullet(doc,
        "A Kolmogorov\u2013Smirnov (KS) test compares the observed p-value distribution from each "
        "null permutation against the expected uniform distribution. Systematic departure from "
        "uniformity in the null suggests residual confounding or model misspecification.",
        bold_prefix="KS uniformity test")
    add_body(doc,
        "The null FPR feeds directly into the CAF Calibration component (Section S3.11). A well-calibrated "
        "pipeline produces null FPR close to the nominal 5% threshold. Values substantially above 5% "
        "indicate inflation; values substantially below indicate over-correction.")

    # S3.14
    add_heading(doc, 'S3.14 Covariate Governance', level=2, font_size=9)
    add_body(doc,
        "The apply_covariate_governance() function enforces a set of rules on the candidate covariate "
        "list before fitting the limma model. This prevents model instability from too many, redundant, "
        "or poorly behaved covariates.")
    add_bullet(doc,
        "Users may specify --force-covariates (always included regardless of other checks), "
        "--allow-covariates (eligible for auto-selection), and --block-covariates (always excluded). "
        "Forced covariates bypass all governance filters.",
        bold_prefix="Forced / allowed / blocked lists")
    add_bullet(doc,
        "Covariates with missingness above a configurable threshold are either dropped or imputed "
        "(median for numeric, mode for categorical). The decision is logged in the decision ledger.",
        bold_prefix="Missingness handling")
    add_bullet(doc,
        "Pairwise Pearson correlations are computed among all numeric candidate covariates. When "
        "|r| \u2265 0.8 between two covariates, the one with weaker association to the top PCs is removed.",
        bold_prefix="Collinearity check")
    add_bullet(doc,
        "To preserve degrees of freedom, the total number of covariates (including surrogate variables) "
        "is capped at floor(n_min / 3), where n_min is the smaller group size. Covariates are ranked "
        "by PC association strength, and lower-ranked covariates are dropped if the cap is exceeded.",
        bold_prefix="Overcorrection guard")


def build_s4(doc):
    """S4. Dashboard Navigation Guide"""
    add_heading(doc, 'S4. Dashboard Navigation Guide', level=1, font_size=11)

    add_body(doc,
        "IlluMeta generates a self-contained interactive HTML dashboard that serves as the primary "
        "interface for exploring results. The dashboard is organized into the following sections:")

    add_heading(doc, 'S4.1 Executive Summary', level=2, font_size=9)
    add_body(doc,
        "The top of the dashboard displays a verdict badge (High/Moderate/Low confidence) based on "
        "method concordance, signal stability, and negative control calibration. Key Performance "
        "Indicators (KPIs) show the primary branch, intersection DMP count, and per-pipeline DMP counts. "
        "A warnings panel highlights critical issues (e.g., high lambda, residual batch effects).")

    add_heading(doc, 'S4.2 Beginner Path (Start Here)', level=2, font_size=9)
    add_body(doc,
        "A 3-step guide for new users:")
    add_bullet(doc, "Check sample quality. Verify QC pass/fail rates and signal quality before interpreting results.",
               bold_prefix="Step 1")
    add_bullet(doc, "Review the High-confidence Intersection. The Intersection (Native) tab reports consensus DMPs "
               "supported by both pipelines while preserving SeSAMe masking. Use this as the recommended primary "
               "consensus set, and compare against Intersection (Strict) as a conservative complete-case cross-check.",
               bold_prefix="Step 2")
    add_bullet(doc, "Dive into individual pipelines. Explore Minfi and Sesame tabs for method-specific depth, "
               "volcano plots, Manhattan plots, and DMR results.",
               bold_prefix="Step 3")

    add_heading(doc, 'S4.3 Run Controls & QC', level=2, font_size=9)
    add_body(doc,
        "Displays the analysis parameters used (thresholds, tissue type, array type, CRF tier) "
        "and a QC summary showing samples passed/failed, probes retained/removed at each filtering step, "
        "and sex mismatch status.")

    add_heading(doc, 'S4.4 Pipeline Tabs', level=2, font_size=9)
    add_body(doc,
        "Results are organized into multiple tabs, selectable via the tabbed interface:")
    add_table(doc,
        ['Tab', 'Contents', 'When to use'],
        [
            ['Intersection (Native)', 'Consensus DMPs, concordance, overlap', 'Primary results (recommended)'],
            ['Intersection (Strict)', 'Strict intersection DMPs', 'Alternative conservative view'],
            ['Minfi (Noob)', 'Full Minfi pipeline results', 'Method-specific exploration'],
            ['Sesame (Strict)', 'Sesame Minfi-aligned results', 'Method-specific exploration'],
            ['Sesame (Native)', 'Sesame pOOBAH-native results', 'Strictest QC masking view'],
        ],
        col_widths=[3.5, 5.0, 5.5])

    add_heading(doc, 'S4.5 Per-tab sections', level=2, font_size=9)
    add_body(doc, "Each pipeline tab contains the following expandable sections:")
    add_bullet(doc, "Volcano plot, Manhattan plot, top 100 heatmap, searchable DMP table.",
               bold_prefix="Primary Results")
    add_bullet(doc, "DMR volcano, DMR Manhattan, top DMRs heatmap, searchable DMR table.",
               bold_prefix="Region-Level Results")
    add_bullet(doc, "PCA before/after correction, sample clustering, Q-Q plot, PVCA before/after.",
               bold_prefix="Quality & Diagnostics")
    add_bullet(doc, "Batch method comparison, batch evaluation before/after, auto-covariate list, dropped covariates.",
               bold_prefix="Batch & Covariates")
    add_bullet(doc, "Epigenetic age predictions (methylclock), placental age (planet, if applicable).",
               bold_prefix="Clocks & Age")


def build_s5(doc):
    """S5. Output File Reference"""
    add_heading(doc, 'S5. Output File Reference', level=1, font_size=11)

    add_body(doc,
        "IlluMeta produces approximately 90-100 output files per analysis. The table below provides "
        "a complete reference organized by category. Pipeline-specific files are prefixed with "
        "Minfi_, Sesame_, or Sesame_Native_.")

    # Config & Metadata
    add_heading(doc, 'S5.1 Configuration & Metadata', level=2, font_size=9)
    add_table(doc,
        ['File', 'Format', 'Description'],
        [
            ['summary.json', 'JSON', 'High-level analysis summary: DMP counts, CAF scores, tier, primary branch'],
            ['analysis_parameters.json', 'JSON', 'Complete analysis configuration and thresholds'],
            ['preflight_report.json', 'JSON', 'Pre-analysis validation: sample counts, group distribution, IDAT checks'],
            ['code_version.txt', 'TXT', 'Git commit hash for reproducibility'],
            ['sessionInfo.txt', 'TXT', 'R version, platform, and loaded packages'],
            ['methods.md', 'MD', 'Auto-generated methods text suitable for paper Methods section'],
            ['decision_ledger.tsv', 'TSV', 'Automated decision log: batch strategy, covariate selection, branch choice'],
            ['config_used.yaml', 'YAML', 'Resolved configuration and preset details'],
        ],
        col_widths=[4.5, 1.0, 9.0])

    # QC
    add_heading(doc, 'S5.2 Quality Control', level=2, font_size=9)
    add_table(doc,
        ['File', 'Format', 'Description'],
        [
            ['Preflight_Summary.csv', 'CSV', 'Pre-analysis checks: sample counts, group sizes, IDAT modal size'],
            ['Preflight_IDAT_Pairs.csv', 'CSV', 'Per-sample IDAT pair existence validation'],
            ['Input_Profile.csv', 'CSV', 'Metadata column profiling: data types, missing values, unique counts'],
            ['Input_Group_Distribution.csv', 'CSV', 'Sample group breakdown (group name and count)'],
            ['QC_Summary.csv', 'CSV', 'QC metrics: samples passed/failed, probes at each filtering step'],
            ['Sample_QC_Metrics.csv', 'CSV', 'Per-sample QC: detection fail fraction, intensity medians'],
            ['Probe_Filter_Summary.csv', 'CSV', 'Step-by-step probe filtering: raw, detectionP, cross-reactive, SNP, sex'],
            ['CrossReactive_Removed_Probes.csv', 'CSV', 'List of removed cross-reactive probes with source'],
        ],
        col_widths=[5.0, 1.0, 8.5])

    # Pipeline results
    add_heading(doc, 'S5.3 Pipeline Results (per pipeline)', level=2, font_size=9)
    add_body(doc, "The following files are produced for each pipeline (Minfi, Sesame, Sesame_Native), "
                  "prefixed accordingly (e.g., Minfi_DMPs_full.csv):")
    add_table(doc,
        ['File suffix', 'Format', 'Description'],
        [
            ['_DMPs_full.csv', 'CSV', 'Full DMP results: CpG, Gene, logFC, P, FDR, Beta values, genomic location'],
            ['_DMRs.csv', 'CSV', 'DMR results: genomic coordinates, CpG count, p-value, gene annotations'],
            ['_Metrics.csv', 'CSV', 'Pipeline metrics: lambda, batch correction method, SV count, signal mode'],
            ['_Signal_Preservation.csv', 'CSV', 'Signal preservation: logFC correlation, sign concordance, Jaccard index'],
            ['_VariancePartition.csv', 'CSV', 'Variance explained by each factor (group, covariates, residuals)'],
            ['_PVCA.csv', 'CSV', 'Principal variance component analysis proportions'],
            ['_AfterCorrection_PVCA.csv', 'CSV', 'PVCA after batch correction (Sesame/Sesame_Native)'],
            ['_AutoCovariates.csv', 'CSV', 'Auto-selected covariate candidates with PC associations and decisions'],
            ['_Batch_Confounding.csv', 'CSV', 'Batch variable confounding analysis: Cramer\'s V, R-squared, tier'],
            ['_Batch_Evaluation_Before/After_Table.csv', 'CSV', 'Covariate-PC associations before/after correction'],
            ['_CAF_Summary.csv', 'CSV', 'CAF component scores: calibration, preservation, batch removal, CAI'],
            ['_CAF_Report.txt', 'TXT', 'Detailed CAF narrative report'],
            ['_CRF_Report.txt', 'TXT', 'CRF robustness assessment report'],
            ['_CRF_SSS_Summary.csv', 'CSV', 'Signal stability score: top-k overlap and sign agreement'],
            ['_DroppedCovariates.csv', 'CSV', 'Covariates excluded from model with reasons'],
            ['_Epigenetic_Age_methylclock.csv', 'CSV', 'Clock predictions: Horvath, Hannum, Levine, etc.'],
            ['_Permutation_Results.csv', 'CSV', 'Per-CpG permutation null test results'],
            ['_Permutation_Summary.csv', 'CSV', 'Permutation summary: mean/median/max significant hits under null'],
            ['_Ablation_Summary.csv', 'CSV', 'Batch metrics before vs. after correction (ablation check)'],
        ],
        col_widths=[5.0, 1.0, 8.5])

    # Intersection
    add_heading(doc, 'S5.4 Intersection & Consensus', level=2, font_size=9)
    add_table(doc,
        ['File', 'Format', 'Description'],
        [
            ['Intersection_Consensus_DMPs.csv', 'CSV', 'Strict consensus DMPs: both pipelines significant + concordant'],
            ['Intersection_Native_Consensus_DMPs.csv', 'CSV', 'Native consensus DMPs (using Sesame Native)'],
            ['Intersection_Discordant_Probes.csv', 'CSV', 'Probes significant in both but with opposite effect direction'],
            ['Intersection_Comparison_Metrics.csv', 'CSV', 'Pipeline overlap statistics and concordance metrics'],
            ['Intersection_Native_Discordant_Probes.csv', 'CSV', 'Discordant probes for Native intersection'],
            ['Intersection_Native_Comparison_Metrics.csv', 'CSV', 'Native comparison metrics'],
        ],
        col_widths=[5.5, 1.0, 8.0])

    # Cell & Clocks
    add_heading(doc, 'S5.5 Cell Deconvolution & Clocks', level=2, font_size=9)
    add_table(doc,
        ['File', 'Format', 'Description'],
        [
            ['Cell_Deconvolution_Summary.csv', 'CSV', 'Cell deconvolution metadata: method, tissue, K, cell types'],
            ['cell_counts_RefFree.csv', 'CSV', 'RefFreeEWAS cell proportions (Minfi): Cell_Latent1-5'],
            ['cell_counts_RefFree_Sesame.csv', 'CSV', 'RefFreeEWAS cell proportions (Sesame): Cell_Latent1-5'],
            ['cell_counts_merged.csv', 'CSV', 'Consolidated Minfi cell counts per sample'],
            ['Cell_Group_Association.csv', 'CSV', 'Cell type association with group: P-value, Eta-squared'],
        ],
        col_widths=[5.0, 1.0, 8.5])

    # CRF & CAF
    add_heading(doc, 'S5.6 Correction Robustness & Adequacy', level=2, font_size=9)
    add_table(doc,
        ['File', 'Format', 'Description'],
        [
            ['CRF_Sample_Tier.csv', 'CSV', 'CRF tier assignment: tier, total_n, min_per_group, warnings'],
            ['CRF_SSS_Summary.csv', 'CSV', 'Global stability scores: top-k overlap and sign agreement'],
            ['Tier3_Eligibility.csv', 'CSV', 'Tier3 meta-analysis eligibility per batch stratum'],
            ['Correction_Adequacy_Summary.csv', 'CSV', 'CAF component scores for primary branch'],
            ['Correction_Adequacy_Report.txt', 'TXT', 'Detailed CAF narrative report'],
            ['Correction_Robustness_Report.txt', 'TXT', 'CRF robustness narrative report'],
        ],
        col_widths=[5.0, 1.0, 8.5])

    # Data matrices
    add_heading(doc, 'S5.7 Data Matrices', level=2, font_size=9)
    add_table(doc,
        ['File', 'Format', 'Description'],
        [
            ['*_BetaMatrix.tsv.gz', 'TSV.GZ', 'Processed beta matrix used for DMP modeling (per pipeline)'],
            ['*_MvalueMatrix.tsv.gz', 'TSV.GZ', 'Processed M-value matrix (logit-transformed betas)'],
            ['*_BetaMatrix_PreFilter.tsv.gz', 'TSV.GZ', 'Pre-filter beta matrix (before sample/probe QC)'],
            ['*_MvalueMatrix_PreFilter.tsv.gz', 'TSV.GZ', 'Pre-filter M-value matrix (before QC)'],
            ['Minfi_DetectionP_PreFilter.tsv.gz', 'TSV.GZ', 'Detection p-value matrix (all samples, all probes)'],
        ],
        col_widths=[5.5, 1.0, 8.0])


def build_s6(doc, showcase_dir=None):
    """S6. Interpreting Results (Beginner Guide)"""
    add_heading(doc, 'S6. Interpreting Results (Beginner Guide)', level=1, font_size=11)

    if showcase_dir:
        gse_name = Path(showcase_dir).parent.name
        add_body(doc,
            "This section explains how to read and interpret the key plots and metrics in IlluMeta's "
            "output. If you are new to DNA methylation analysis, start here. Example figures are shown "
            f"from the Minfi pipeline of {gse_name}, selected as the showcase dataset from our benchmark.")
    else:
        add_body(doc,
            "This section explains how to read and interpret the key plots and metrics in IlluMeta's "
            "output. If you are new to DNA methylation analysis, start here.")

    # S6.1 Volcano
    add_heading(doc, 'S6.1 How to read a Volcano plot', level=2, font_size=9)
    add_body(doc,
        "A volcano plot shows the relationship between effect size (x-axis: log-fold change or delta-beta) "
        "and statistical significance (y-axis: \u2212log\u2081\u2080 adjusted p-value). Each dot represents one CpG site.")
    add_bullet(doc, "CpGs that are both highly significant (high y-axis) and have large effect sizes "
               "(far from center on x-axis) are the most biologically meaningful.",
               bold_prefix="Key interpretation")
    add_bullet(doc, "Red/colored dots are significant (FDR < 0.05); gray dots are not significant.",
               bold_prefix="Color coding")
    add_bullet(doc, "Left side = hypomethylated in test vs. control; right side = hypermethylated in test.",
               bold_prefix="Direction")
    if showcase_dir:
        fig = os.path.join(showcase_dir, 'Minfi_Volcano.png')
        add_figure(doc, fig,
            f"Figure S3. Volcano plot (Minfi pipeline, {gse_name}). Each point is a CpG site; "
            "colored points pass FDR < 0.05.", width=5.0)

    # S6.2 Manhattan
    add_heading(doc, 'S6.2 How to read a Manhattan plot', level=2, font_size=9)
    add_body(doc,
        "A Manhattan plot displays the genomic location of each CpG (x-axis: chromosomal position) "
        "versus its significance (y-axis: \u2212log\u2081\u2080 p-value). Alternating colors distinguish chromosomes.")
    add_bullet(doc, "Look for 'peaks' (clusters of significant CpGs in the same genomic region). "
               "These suggest coordinated methylation changes and are often more biologically meaningful "
               "than isolated hits.",
               bold_prefix="What to look for")
    add_bullet(doc, "The horizontal red line indicates the genome-wide significance threshold.",
               bold_prefix="Significance line")
    if showcase_dir:
        fig = os.path.join(showcase_dir, 'Minfi_Manhattan.png')
        add_figure(doc, fig,
            f"Figure S4. Manhattan plot (Minfi pipeline, {gse_name}). Peaks indicate genomic regions "
            "with clusters of differentially methylated CpGs.", width=5.5)

    # S6.3 QQ
    add_heading(doc, 'S6.3 How to read a Q-Q plot and lambda', level=2, font_size=9)
    add_body(doc,
        "A quantile-quantile (Q-Q) plot compares observed p-values against the expected uniform "
        "distribution. The genomic inflation factor (lambda) summarizes overall inflation:")
    add_bullet(doc, "Points fall along the diagonal. No systematic bias.",
               bold_prefix="Lambda near 1.0 (ideal)")
    add_bullet(doc, "Points curve upward from the diagonal. May indicate batch effects, population "
               "stratification, or genuine widespread signal. Values > 1.5 warrant investigation.",
               bold_prefix="Lambda > 1.2 (inflated)")
    add_bullet(doc, "Points fall below the diagonal. May indicate over-correction. Check if batch "
               "correction removed too much signal.",
               bold_prefix="Lambda < 0.9 (deflated)")
    if showcase_dir:
        fig = os.path.join(showcase_dir, 'Minfi_QQPlot.png')
        add_figure(doc, fig,
            f"Figure S5. Q-Q plot (Minfi pipeline, {gse_name}). Deviation from the diagonal indicates "
            "inflation or deflation of test statistics.", width=4.0)

    # S6.4 PCA
    add_heading(doc, 'S6.4 Understanding PCA plots', level=2, font_size=9)
    add_body(doc,
        "PCA (Principal Component Analysis) reduces the complexity of methylation data into a few "
        "dimensions. IlluMeta shows PCA before and after batch correction:")
    add_bullet(doc, "Samples from different groups should separate. If control and test samples overlap "
               "completely, the methylation difference between groups may be small.",
               bold_prefix="Before correction")
    add_bullet(doc, "Batch-driven clustering should disappear while group separation is preserved. "
               "If groups merge after correction, the correction may have been too aggressive.",
               bold_prefix="After correction")
    if showcase_dir:
        fig_before = os.path.join(showcase_dir, 'Minfi_PCA_Before.png')
        fig_after = os.path.join(showcase_dir, 'Minfi_PCA_After_Correction.png')
        add_figure(doc, fig_before,
            f"Figure S6a. PCA before batch correction (Minfi pipeline, {gse_name}). "
            "Sample colors indicate group membership.", width=4.0)
        add_figure(doc, fig_after,
            f"Figure S6b. PCA after batch correction (Minfi pipeline, {gse_name}). "
            "Batch-driven clustering should be reduced while group separation is preserved.", width=4.0)

    # S6.5 PVCA
    add_heading(doc, 'S6.5 Understanding PVCA', level=2, font_size=9)
    add_body(doc,
        "PVCA (Principal Variance Component Analysis) shows what proportion of total methylation "
        "variance is explained by each factor (group, batch, covariates, residuals). After correction, "
        "the primary_group proportion should increase and batch-related proportions should decrease.")
    if showcase_dir:
        fig = os.path.join(showcase_dir, 'Minfi_PVCA.png')
        add_figure(doc, fig,
            f"Figure S7. PVCA variance decomposition (Minfi pipeline, {gse_name}). "
            "Each bar shows the proportion of variance attributed to a specific factor.", width=4.0)

    # S6.6 Top 100 heatmap
    add_heading(doc, 'S6.6 How to read the Top 100 DMP heatmap', level=2, font_size=9)
    add_body(doc,
        "The top 100 DMP heatmap shows the beta values (methylation levels) of the 100 most significant "
        "CpGs across all samples. Rows are CpGs (clustered by similarity) and columns are samples "
        "(grouped by condition). A clear block pattern\u2014where control and test samples show distinct "
        "methylation profiles\u2014indicates robust differential methylation.")
    if showcase_dir:
        fig = os.path.join(showcase_dir, 'Minfi_Top100_Heatmap.png')
        add_figure(doc, fig,
            f"Figure S8. Top 100 DMP heatmap (Minfi pipeline, {gse_name}). Rows are CpGs; columns are "
            "samples grouped by condition. Red = high methylation, blue = low methylation.", width=5.0)

    # S6.7 DMR analysis
    add_heading(doc, 'S6.7 How to read DMR results', level=2, font_size=9)
    add_body(doc,
        "Differentially methylated regions (DMRs) are genomic intervals where multiple adjacent CpGs "
        "show coordinated methylation changes. The DMR volcano plot displays each region as a point, "
        "with the x-axis showing the mean effect size across the region and the y-axis showing "
        "significance. DMRs provide stronger evidence of regulatory impact than individual CpGs because "
        "coordinated changes across a region are less likely to arise from noise.")
    if showcase_dir:
        fig = os.path.join(showcase_dir, 'Minfi_DMR_Volcano.png')
        add_figure(doc, fig,
            f"Figure S9. DMR volcano plot (Minfi pipeline, {gse_name}). Each point represents a "
            "differentially methylated region identified by dmrff.", width=5.0)

    # S6.8 Sample clustering
    add_heading(doc, 'S6.8 Sample clustering dendrogram', level=2, font_size=9)
    add_body(doc,
        "The sample clustering dendrogram shows hierarchical clustering of samples based on their "
        "genome-wide methylation profiles (Euclidean distance on beta values). Samples from the same "
        "group should cluster together. If samples cluster by batch rather than group, this may indicate "
        "unresolved batch effects.")
    if showcase_dir:
        fig = os.path.join(showcase_dir, 'Minfi_Sample_Clustering_Distance.png')
        add_figure(doc, fig,
            f"Figure S10. Sample clustering dendrogram (Minfi pipeline, {gse_name}). "
            "Branch colors indicate group membership.", width=4.5)

    # S6.9 Batch evaluation
    add_heading(doc, 'S6.9 Batch evaluation heatmaps', level=2, font_size=9)
    add_body(doc,
        "Batch evaluation heatmaps show the statistical association between known covariates/batch "
        "variables and the top principal components of the methylation data. Cells are colored by "
        "\u2212log\u2081\u2080(p-value): red cells indicate strong associations. Comparing before and after "
        "correction, batch-related associations should diminish while group associations remain.")
    if showcase_dir:
        fig_before = os.path.join(showcase_dir, 'Minfi_Batch_Evaluation_Before.png')
        fig_after = os.path.join(showcase_dir, 'Minfi_Batch_Evaluation_After.png')
        add_figure(doc, fig_before,
            f"Figure S11a. Batch evaluation heatmap before correction (Minfi pipeline, {gse_name}). "
            "Red cells indicate strong covariate\u2013PC associations.", width=4.5)
        add_figure(doc, fig_after,
            f"Figure S11b. Batch evaluation heatmap after correction (Minfi pipeline, {gse_name}). "
            "Batch-associated cells should fade while group signal persists.", width=4.5)

    # S6.10 Intersection plots
    add_heading(doc, 'S6.10 Intersection overlap and concordance', level=2, font_size=9)
    add_body(doc,
        "The intersection plots show the overlap between the two pipelines. The significant overlap "
        "Venn diagram displays how many DMPs are unique to Minfi, unique to SeSAMe, or shared (consensus). "
        "The logFC concordance scatter plot shows the correlation of effect sizes between pipelines for "
        "shared CpGs. High concordance (points near the diagonal) confirms that both methods detect "
        "similar biological signals.")
    if showcase_dir:
        fig_overlap = os.path.join(showcase_dir, 'Intersection_Significant_Overlap.png')
        fig_concordance = os.path.join(showcase_dir, 'Intersection_LogFC_Concordance.png')
        add_figure(doc, fig_overlap,
            f"Figure S12a. Pipeline overlap ({gse_name}). Venn diagram of significant DMPs "
            "from Minfi and SeSAMe pipelines.", width=3.5)
        add_figure(doc, fig_concordance,
            f"Figure S12b. LogFC concordance ({gse_name}). Scatter plot of per-CpG log-fold "
            "changes between Minfi and SeSAMe for shared significant CpGs.", width=4.0)

    # S6.11 Consensus DMPs
    add_heading(doc, 'S6.11 What consensus DMPs mean', level=2, font_size=9)
    add_body(doc,
        "Consensus DMPs are CpG sites that are statistically significant in both the Minfi and SeSAMe "
        "pipelines with the same direction of effect. Because two independent methods agree, these hits "
        "are more likely to be genuine biological signals rather than pipeline-specific artifacts. "
        "The consensus approach is deliberately conservative: it may miss some true signals that only "
        "one method detects, but the signals it does report have high confidence.")

    # S6.12 CRF tier
    add_heading(doc, 'S6.12 What CRF tier means for your study', level=2, font_size=9)
    add_body(doc,
        "The CRF tier reflects your study's statistical power based on sample size:")
    add_bullet(doc, "n < 12. Results are exploratory only. SVA and advanced diagnostics are disabled. "
               "Plan for replication in a larger cohort.",
               bold_prefix="Minimal")
    add_bullet(doc, "n = 12\u201323. Independent validation is required. Limited batch correction options.",
               bold_prefix="Small")
    add_bullet(doc, "n = 24\u201349. Full pipeline available with reasonable statistical power.",
               bold_prefix="Moderate")
    add_bullet(doc, "n \u2265 50. All robustness checks are fully powered. Best for publication.",
               bold_prefix="Large")

    # S6.13 Batch effects
    add_heading(doc, 'S6.13 When to worry about batch effects', level=2, font_size=9)
    add_body(doc, "You should investigate batch effects when:")
    add_bullet(doc, "PCA 'before' plot shows samples clustering by batch (e.g., Sentrix slide) rather than by group.")
    add_bullet(doc, "Lambda is much greater than 1.2 after correction.")
    add_bullet(doc, "The batch evaluation heatmap shows many significant associations between batch variables and PCs after correction.")
    add_bullet(doc, "The CAF batch removal score is low (< 0.5).")
    add_body(doc,
        "If batch effects persist after automatic correction, consider: (1) checking for sample swaps, "
        "(2) adding the batch variable to the design explicitly, or (3) performing a stratified/Tier3 analysis.")


def build_s7(doc, benchmark_data, figures_dir):
    """S7. Benchmark Results"""
    add_heading(doc, 'S7. Benchmark Results', level=1, font_size=11)

    platform_order = {'450k': 0, 'EPIC': 1, 'EPICv2': 2}
    n_datasets = len(benchmark_data)
    total_samples = sum(safe_int(d.get('n_samples', 0)) for d in benchmark_data)
    platforms_list = sorted(set(d.get('platform', '') for d in benchmark_data if d.get('platform')),
                            key=lambda p: platform_order.get(p, 9))
    platforms_disp = []
    for p in platforms_list:
        if str(p).lower() == "450k":
            platforms_disp.append("450K")
        elif str(p) == "EPICv2":
            platforms_disp.append("EPIC v2")
        else:
            platforms_disp.append(str(p))
    if len(platforms_disp) > 2:
        platform_disp_str = ", ".join(platforms_disp[:-1]) + ", and " + platforms_disp[-1]
    else:
        platform_disp_str = " and ".join(platforms_disp)
    sample_range = sorted([safe_int(d.get('n_samples', 0)) for d in benchmark_data]) if benchmark_data else [0]

    add_body(doc,
        f"IlluMeta was benchmarked on {n_datasets} publicly available GEO datasets "
        f"(total n = {total_samples}; range {sample_range[0]}\u2013{sample_range[-1]}) spanning "
        f"{platform_disp_str} arrays.")

    # Table S1
    add_heading(doc, 'Table S1. Full benchmark summary', level=2, font_size=9)
    # Sort by platform then sample size (matching Figure 2 / Table 2 order)
    benchmark_data = sorted(benchmark_data,
                            key=lambda r: (platform_order.get(r.get('platform', ''), 9),
                                           safe_int(r.get('n_samples', 0))))
    rows = []
    for d in benchmark_data:
        rows.append([
            d.get('gse_id', ''),
            d.get('platform', ''),
            str(d.get('n_samples', '')),
            d.get('crf_tier', '').capitalize(),
            f"{safe_int(d.get('minfi_dmps', 0)):,}",
            f"{safe_int(d.get('sesame_dmps', 0)):,}",
            f"{safe_int(d.get('consensus_dmps', 0)):,}",
            f"{safe_int(d.get('minfi_dmps_nominal', 0)):,}",
            f"{safe_int(d.get('sesame_dmps_nominal', 0)):,}",
            f"{safe_int(d.get('minfi_dmrs', 0)):,}",
            f"{safe_int(d.get('sesame_dmrs', 0)):,}",
            f"{safe_float(d.get('lambda_minfi', 1.0)):.3f}",
            f"{safe_float(d.get('lambda_sesame', 1.0)):.3f}",
            f"{safe_float(d.get('caf_score', 0)):.3f}",
        ])
    add_table(doc,
        ['Dataset', 'Platform', 'n', 'CRF Tier',
         'Minfi DMPs', 'Sesame DMPs', 'Consensus DMPs',
         'Minfi DMPs*', 'Sesame DMPs*',
         'Minfi DMRs', 'Sesame DMRs',
         '\u03bb (Minfi)', '\u03bb (Sesame)', 'CAF'],
        rows,
        col_widths=[1.4, 0.8, 0.5, 1.0, 1.0, 1.0, 1.2, 1.0, 1.0, 1.0, 1.0, 0.9, 0.9, 0.7])
    add_body(doc, 'DMPs: FDR < 0.05. DMPs*: nominal (raw P < 0.05, |log\u2082FC| > 1). '
             'DMRs: dmrff regions (FDR < 0.05). '
             '\u03bb: genomic inflation factor (diagnostic; may increase with true signal). '
             'CAF: Correction Adequacy score.',
             font_size=7, italic=True)

    # Figure S2
    fig2 = os.path.join(figures_dir, 'FigureS2_Ablation_Lambda.png')
    add_figure(doc, fig2,
        f"Figure S2. Effect of batch correction on genomic inflation (lambda) across {n_datasets} GEO datasets. "
        "Raw vs corrected lambda is shown for Minfi and SeSAMe.", width=5.5)


def build_s8(doc):
    """S8. Command-Line Reference"""
    add_heading(doc, 'S8. Command-Line Reference', level=1, font_size=11)

    # download
    add_heading(doc, 'S8.1 illumeta.py download', level=2, font_size=9)
    add_body(doc, "Download IDAT files and metadata from GEO.")
    add_code_block(doc, "python3 illumeta.py download <GSE_ID> [options]")
    add_table(doc,
        ['Option', 'Default', 'Description'],
        [
            ['GSE_ID', '(required)', 'GEO Series accession (e.g., GSE121633)'],
            ['-o, --out-dir', './<GSE_ID>', 'Output directory for downloaded files'],
            ['--platform', '(auto)', 'Force a specific GPL platform (e.g., GPL21145)'],
        ],
        col_widths=[3.5, 2.0, 9.0])

    # analysis
    add_heading(doc, 'S8.2 illumeta.py analysis', level=2, font_size=9)
    add_body(doc, "Run the full dual-pipeline analysis.")
    add_code_block(doc, "python3 illumeta.py analysis -i <input_dir> --group_con <con> --group_test <test> [options]")
    add_table(doc,
        ['Option', 'Default', 'Description'],
        [
            ['-i, --input-dir', '(required)', 'Input directory containing IDATs and configure.tsv'],
            ['--group_con', '(required)', 'Control group label'],
            ['--group_test', '(required)', 'Test group label'],
            ['--pval', '0.05', 'FDR threshold for significance'],
            ['--lfc', '0.5', 'Minimum absolute log-fold change'],
            ['--delta-beta', '0', 'Minimum absolute delta-beta'],
            ['--tissue', 'Auto', 'Tissue type for cell deconvolution (Auto, Blood, CordBlood, Placenta, DLPFC)'],
            ['--auto-group', 'off', 'Auto-populate primary_group from metadata'],
            ['--group-column', 'None', 'Metadata column for auto-grouping'],
            ['--group-key', 'None', 'GEO characteristics key for auto-grouping'],
            ['--group-map', 'None', 'Value mapping for auto-grouping (e.g., "normal=Control,tumor=Case")'],
            ['--skip-sesame', 'off', 'Skip SeSAMe pipeline (Minfi only)'],
            ['--disable-sva', 'off', 'Disable Surrogate Variable Analysis'],
            ['--include-covariates', 'None', 'Comma-separated list of covariates to include'],
            ['--include-clock-covariates', 'off', 'Use clock estimates as candidate covariates'],
            ['--permutations', '20', 'Number of label permutations for null calibration'],
            ['--beginner-safe', 'off', 'Enable conservative thresholds and stricter checks'],
            ['--sesame-typeinorm', 'off', 'Enable SeSAMe dye-bias correction TypeINorm'],
            ['--marker-list', 'None', 'CpG marker list for signal preservation check'],
            ['--cell-reference', 'None', 'Custom cell reference (package, object, or .rds/.rda)'],
            ['--output', 'auto', 'Custom output directory'],
        ],
        col_widths=[3.5, 1.5, 9.5])

    # search
    add_heading(doc, 'S8.3 illumeta.py search', level=2, font_size=9)
    add_body(doc, "Search GEO for IDAT-enabled methylation datasets.")
    add_code_block(doc, 'python3 illumeta.py search --keywords "breast cancer" --email user@example.com')
    add_table(doc,
        ['Option', 'Default', 'Description'],
        [
            ['-k, --keywords', '(required)', 'Search keywords (use quotes for multi-word)'],
            ['-o, --output', 'geo_idat_methylation.tsv', 'Output TSV path'],
            ['--email', 'None', 'Email for NCBI Entrez API'],
            ['--no-check-suppl', 'off', 'Skip FTP supplement checks (faster)'],
        ],
        col_widths=[3.5, 3.5, 7.5])

    # doctor
    add_heading(doc, 'S8.4 illumeta.py doctor', level=2, font_size=9)
    add_body(doc, "Verify that all required R packages and dependencies are installed correctly.")
    add_code_block(doc, "python3 illumeta.py doctor")
    add_body(doc,
        "The doctor command checks: core R packages, optional packages (clocks, EPIC v2), "
        "Python dependencies, and system tools (pandoc). It reports OK/missing status for each component.")


def build_s9(doc):
    """S9. FAQ & Troubleshooting"""
    add_heading(doc, 'S9. FAQ & Troubleshooting', level=1, font_size=11)

    add_heading(doc, 'S9.1 Common errors and solutions', level=2, font_size=9)
    add_table(doc,
        ['Error', 'Cause', 'Solution'],
        [
            ['xml2/libxml-2.0 not found', 'Missing system library', 'Install libxml2-dev (Ubuntu) or brew install libxml2 (macOS)'],
            ['lzma.h not found', 'Missing compression library', 'Install liblzma-dev (Ubuntu) or brew install xz (macOS)'],
            ['pandoc: command not found', 'pandoc not installed', 'sudo apt install pandoc (Ubuntu) or brew install pandoc (macOS)'],
            ['pthread_create error', 'Sesame threading issue', 'Set SESAME_NTHREAD=1 before running analysis'],
            ['No finite residual std devs', 'Very small n or zero variance', 'Reduce covariates or disable batch correction for tiny cohorts'],
            ['Contrasts require 2+ levels', 'Single-level covariate after NA filtering', 'Remove or merge constant columns in configure.tsv'],
            ['Package built under R x.y', 'R version mismatch', 'Rerun: ILLUMETA_CLEAN_MISMATCHED_RLIB=1 Rscript r_scripts/setup_env.R'],
            ['GEO download NULL/404', 'GEO FTP issue', 'Manually download GSE*_RAW.tar and place under projects/<GSE>/<GSE>/'],
        ],
        col_widths=[3.5, 3.0, 8.0])

    add_heading(doc, 'S9.2 Sample size recommendations', level=2, font_size=9)
    add_table(doc,
        ['Total samples', 'CRF tier', 'Expected results', 'Recommendation'],
        [
            ['< 12', 'Minimal', 'Exploratory; very limited power', 'Use for pilot studies; plan replication'],
            ['12\u201323', 'Small', 'Can detect large effects', 'Validate findings independently'],
            ['24\u201349', 'Moderate', 'Reasonable power', 'Suitable for focused EWAS studies'],
            ['\u2265 50', 'Large', 'Good power; robust analysis', 'Publication-ready with comprehensive QC'],
        ],
        col_widths=[2.0, 1.8, 4.5, 6.0])

    add_heading(doc, 'S9.3 Which results to cite in papers', level=2, font_size=9)
    add_body(doc, "When reporting IlluMeta results in a manuscript, we recommend:")
    add_bullet(doc, "Report Intersection (Native) consensus DMP counts as your primary result.",
               bold_prefix="Primary results")
    add_bullet(doc, "Mention Minfi and Sesame DMP counts for comparison and transparency.",
               bold_prefix="Pipeline comparison")
    add_bullet(doc, "Include the methods.md text (or adapted version) in your Methods section.",
               bold_prefix="Methods text")
    add_bullet(doc, "Include analysis_parameters.json, decision_ledger.tsv, and CRF_Sample_Tier.csv.",
               bold_prefix="Supplementary files")
    add_bullet(doc, "Report lambda, CAF score, and CRF tier for quality assessment.",
               bold_prefix="Quality metrics")
    add_body(doc,
        "Example citation: \"DNA methylation analysis was performed using IlluMeta (v1.0; "
        "https://github.com/kangk1204/illumeta), which runs dual normalization pipelines "
        "(minfi Noob and SeSAMe pOOBAH) and reports consensus DMPs at FDR < 0.05.\"")


def build_s10(doc):
    """S10. References"""
    add_heading(doc, 'S10. References', level=1, font_size=11)

    references = [
        ("Barrett,T. et al. (2013) NCBI GEO: archive for functional genomics data "
         "sets\u2014update. Nucleic Acids Res., 41, D991\u2013D995.", "PMID: 23193258"),
        ("Aryee,M.J. et al. (2014) Minfi: a flexible and comprehensive Bioconductor "
         "package for the analysis of Infinium DNA methylation microarrays. "
         "Bioinformatics, 30, 1363\u20131369.", "PMID: 24478339"),
        ("Morris,T.J. et al. (2014) ChAMP: 450k Chip Analysis Methylation Pipeline. "
         "Bioinformatics, 30, 428\u2013430.", "PMID: 24336642"),
        ("Leek,J.T. et al. (2010) Tackling the widespread and critical impact of batch "
         "effects in high-throughput data. Nat. Rev. Genet., 11, 733\u2013739.", "PMID: 20838408"),
        ("Leek,J.T. and Storey,J.D. (2007) Capturing heterogeneity in gene expression "
         "studies by surrogate variable analysis. PLoS Genet., 3, e161.", "PMID: 17907809"),
        ("Johnson,W.E. et al. (2007) Adjusting batch effects in microarray expression "
         "data using empirical Bayes methods. Biostatistics, 8, 118\u2013127.", "PMID: 16632515"),
        ("Zhou,W. et al. (2018) SeSAMe: reducing artifactual detection of DNA methylation "
         "by Infinium BeadChips in genomic deletions. Nucleic Acids Res., 46, e123.", "PMID: 30085201"),
        ("Pidsley,R. et al. (2016) Critical evaluation of the Illumina MethylationEPIC "
         "BeadChip microarray for whole-genome DNA methylation profiling. Genome Biol., 17, 208.", "PMID: 27717381"),
        ("McCartney,D.L. et al. (2016) Identification of polymorphic and off-target probe "
         "binding sites on the Illumina Infinium MethylationEPIC BeadChip. Genomics Data, 9, 22\u201324.", "PMID: 27330998"),
        ("Ritchie,M.E. et al. (2015) limma powers differential expression analyses for "
         "RNA-sequencing and microarray studies. Nucleic Acids Res., 43, e47.", "PMID: 25605792"),
        ("Teschendorff,A.E. et al. (2017) A comparison of reference-based algorithms "
         "for correcting cell-type heterogeneity in Epigenome-Wide Association Studies. "
         "BMC Bioinformatics, 18, 105.", "PMID: 28193155"),
        ("Houseman,E.A. et al. (2014) Reference-free deconvolution of DNA methylation data "
         "and mediation by cell composition effects. BMC Bioinformatics, 15, 407.", "PMID: 25511081"),
        ("Pelegri-Siso,D. et al. (2021) methylclock: a Bioconductor package to estimate "
         "DNA methylation age. Bioinformatics, 37, 1759\u20131760.", "PMID: 32960939"),
        ("Yuan,V. et al. (2021) Accurate ethnicity prediction from placental DNA methylation "
         "data. Epigenetics Chromatin, 14, 49.", "PMID: 34663446"),
        ("Peters,T.J. et al. (2015) De novo identification of differentially methylated "
         "regions in the human genome. Epigenetics Chromatin, 8, 6.", "PMID: 25972926"),
        ("Suderman,M. et al. (2018) dmrff: identifying differentially methylated "
         "regions efficiently with power and control. bioRxiv, 508556.", "doi: 10.1101/508556"),
    ]

    for i, (ref_text, pmid) in enumerate(references, 1):
        ref_para = doc.add_paragraph()
        ref_para.paragraph_format.space_after = Pt(1)
        ref_para.paragraph_format.space_before = Pt(0)
        ref_para.paragraph_format.line_spacing = Pt(10)

        num_run = ref_para.add_run(f'[{i}] ')
        num_run.font.size = Pt(7.5)
        num_run.font.name = 'Arial'
        num_run.bold = True

        text_run = ref_para.add_run(ref_text + ' ')
        text_run.font.size = Pt(7.5)
        text_run.font.name = 'Arial'

        pmid_run = ref_para.add_run(pmid)
        pmid_run.font.size = Pt(7.5)
        pmid_run.font.name = 'Arial'
        pmid_run.bold = True
        pmid_run.font.color.rgb = RGBColor(0, 0, 180)


# ── Main builder ──────────────────────────────────────────────────────────────

def build_docx(benchmark_data, figures_dir, out_path, showcase_dir=None):
    """Build the complete Supplementary Data DOCX."""
    doc = Document()

    # Page setup: A4, 25.4mm margins
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('Supplementary Data')
    title_run.font.size = Pt(14)
    title_run.font.name = 'Arial'
    title_run.bold = True
    title_para.paragraph_format.space_after = Pt(4)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run(
        'IlluMeta: an automated dual-pipeline framework for robust DNA methylation analysis'
    )
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.name = 'Arial'
    subtitle_run.italic = True
    subtitle_para.paragraph_format.space_after = Pt(2)

    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_para.add_run('Keunsoo Kang')
    author_run.font.size = Pt(10)
    author_run.font.name = 'Arial'
    author_para.paragraph_format.space_after = Pt(2)

    affil_para = doc.add_paragraph()
    affil_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affil_run = affil_para.add_run(
        'Department of Microbiology, Dankook University, Cheonan 31116, Republic of Korea'
    )
    affil_run.font.size = Pt(8)
    affil_run.font.name = 'Arial'
    affil_run.italic = True
    affil_para.paragraph_format.space_after = Pt(12)

    # Table of contents (manual)
    add_heading(doc, 'Contents', level=1, font_size=10)
    toc_items = [
        'S1. Introduction & Overview',
        'S2. Getting Started',
        'S3. Analysis Pipeline (Detailed Methods)',
        '    S3.1 Quality Control  \u2022  S3.2 Normalization  \u2022  S3.3 Batch Detection',
        '    S3.4 Covariate Detection  \u2022  S3.5 DMP Analysis  \u2022  S3.6 DMR Analysis',
        '    S3.7 Consensus Intersection  \u2022  S3.8 Cell Deconvolution  \u2022  S3.9 Clocks',
        '    S3.10 CRF  \u2022  S3.11 CAF Score  \u2022  S3.12 Tier3 Meta-Analysis',
        '    S3.13 Permutation Testing  \u2022  S3.14 Covariate Governance',
        'S4. Dashboard Navigation Guide',
        'S5. Output File Reference',
        'S6. Interpreting Results (with example figures)',
        'S7. Benchmark Results',
        'S8. Command-Line Reference',
        'S9. FAQ & Troubleshooting',
        'S10. References',
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        r = p.add_run(item)
        r.font.size = Pt(9)
        r.font.name = 'Arial'
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(0)
    doc.add_paragraph()  # spacing

    # Build sections
    build_s1(doc, figures_dir)
    build_s2(doc)
    build_s3(doc)
    build_s4(doc)
    build_s5(doc)
    build_s6(doc, showcase_dir=showcase_dir)
    build_s7(doc, benchmark_data, figures_dir)
    build_s8(doc)
    build_s9(doc)
    build_s10(doc)

    # Save
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    print(f"\nSaved Supplementary Data: {out_path}")
    file_size = os.path.getsize(out_path)
    print(f"File size: {file_size:,} bytes")

    # Estimate page count (rough: ~3500 chars per page for A4 9pt)
    # Count paragraphs as proxy
    n_paras = len(doc.paragraphs)
    n_tables = len(doc.tables)
    est_pages = max(1, n_paras // 8 + n_tables * 2)
    print(f"Estimated pages: ~{est_pages}")


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description='Build Supplementary Data DOCX for IlluMeta')
    parser.add_argument('--results-dir', default='benchmarks/application_note_submission_results',
                        help='Directory containing IlluMeta benchmark results')
    parser.add_argument('--figures-dir', default='benchmarks/paper_figures',
                        help='Directory containing generated figures')
    parser.add_argument('--out', default='benchmarks/application_note/IlluMeta_Supplementary_Data.docx',
                        help='Output DOCX path')
    parser.add_argument('--showcase-gse', default=None,
                        help='GSE accession to use as showcase dataset for S6 figures (e.g., GSE141338). Auto-selects highest-CAF dataset if omitted.')
    parser.add_argument('--allow-fallback', action='store_true',
                        help='Allow building a DOCX with hardcoded placeholder benchmark values when results are missing (NOT for submission).')
    args = parser.parse_args()

    print("IlluMeta Supplementary Data DOCX Builder")
    print("=" * 50)

    # Collect benchmark data
    print("\nCollecting benchmark data...")
    benchmark_data = collect_benchmark_data(args.results_dir)
    print(f"Found {len(benchmark_data)} datasets")

    if not benchmark_data:
        if not args.allow_fallback:
            sys.exit(
                "ERROR: No benchmark data found under --results-dir. "
                "Provide a valid results directory (with GSE*/.../summary.json), "
                "or pass --allow-fallback to generate a non-submission placeholder DOCX."
            )
        print("WARNING: No benchmark data found. Using hardcoded fallback (--allow-fallback).")
        benchmark_data = [
            {'gse_id': 'GSE113687', 'platform': 'EPIC', 'n_samples': '20',
             'crf_tier': 'minimal', 'minfi_dmps': 205948, 'sesame_dmps': 193633,
             'consensus_dmps': 171391, 'minfi_dmrs': 14537, 'sesame_dmrs': 13979,
             'lambda_minfi': 2.013, 'lambda_sesame': 1.145, 'caf_score': 0.6001},
            {'gse_id': 'GSE66313', 'platform': '450k', 'n_samples': '29',
             'crf_tier': 'moderate', 'minfi_dmps': 0, 'sesame_dmps': 0,
             'consensus_dmps': 0, 'minfi_dmrs': 2321, 'sesame_dmrs': 1741,
             'lambda_minfi': 1.063, 'lambda_sesame': 1.099, 'caf_score': 0.5051},
            {'gse_id': 'GSE74071', 'platform': '450k', 'n_samples': '6',
             'crf_tier': 'minimal', 'minfi_dmps': 0, 'sesame_dmps': 0,
             'consensus_dmps': 0, 'minfi_dmrs': 3956, 'sesame_dmrs': 1234,
             'lambda_minfi': 6.598, 'lambda_sesame': 4.929, 'caf_score': 0.4670},
            {'gse_id': 'GSE308134', 'platform': 'EPIC', 'n_samples': '25',
             'crf_tier': 'moderate', 'minfi_dmps': 1400, 'sesame_dmps': 1493,
             'consensus_dmps': 807, 'minfi_dmrs': 2394, 'sesame_dmrs': 3384,
             'lambda_minfi': 0.463, 'lambda_sesame': 1.601, 'caf_score': 0.5957},
            {'gse_id': 'GSE281691', 'platform': 'EPIC', 'n_samples': '160',
             'crf_tier': 'large', 'minfi_dmps': 0, 'sesame_dmps': 0,
             'consensus_dmps': 0, 'minfi_dmrs': 2205, 'sesame_dmrs': 874,
             'lambda_minfi': 0.959, 'lambda_sesame': 0.900, 'caf_score': 0.4673},
            {'gse_id': 'GSE312914', 'platform': 'EPICv2', 'n_samples': '147',
             'crf_tier': 'large', 'minfi_dmps': 312, 'sesame_dmps': 314,
             'consensus_dmps': 306, 'minfi_dmrs': 4486, 'sesame_dmrs': 4453,
             'lambda_minfi': 0.952, 'lambda_sesame': 0.888, 'caf_score': 0.4335},
        ]

    # Find showcase dataset for S6 figure embedding
    print("\nFinding showcase dataset...")
    showcase_dir = find_showcase_dataset(args.results_dir, preferred_gse=args.showcase_gse)
    if showcase_dir is None:
        print("WARNING: No showcase dataset found. S6 figures will be text-only.")

    # Build DOCX
    print("\nBuilding DOCX...")
    build_docx(benchmark_data, args.figures_dir, args.out, showcase_dir=showcase_dir)
    print("\nDone!")


if __name__ == '__main__':
    main()
